from ast import If
from math import pi
from PIL import Image
import streamlit as st
from streamlit import session_state as ss
import pandas as pd
import numpy as np
import docx
from docx import Document
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import sys
##punch_dir = (os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
##+ '/pages/punch/')
##sys.path.append(punch_dir)

from punch_word_func import *
from punch_draw_func import *
from punch_text_func import *
from punch_solve_func import *

#Доступные диаметры
dias = [6, 8, 10, 12, 14, 16, 18, 20, 22, 25]
#Параметры поперечного армирования по умолчанию
Rsw = 1.734
Rsw0 = 170
rtype = 'A500'
dsw = 6
sw = 6.0
nsw = 2
Asw = 0.565
qsw = round(Asw*Rsw/sw, 5)
qsw0 = round(Asw*Rsw/sw, 5)
kh0 = 1.5
table_concretes_data = pd.read_excel('RC_data.xlsx', sheet_name="Concretes_SP63", header=[0])
table_reinf_data = pd.read_excel('RC_data.xlsx', sheet_name="Reinforcement_SP63", header=[0])
available_concretes = table_concretes_data['Class'].to_list()
available_reinf = table_reinf_data['Class'].to_list()


st.header('Расчет на продавливание плиты')

is_report = st.sidebar.toggle('Формировать отчет', value=False)

is_init_help, is_geom_full = False, False
if is_report: #Если включена генерация отчета
    is_init_help = st.sidebar.toggle('Расчетные предпосылки', value=False)
    is_geom_full = st.sidebar.toggle('Геом. характеристики подробно', value=False)

#Вывод справки по исходным данным
init_data_help()


if True: #Ввод исходных данных
    cols = st.columns([1, 0.5])
    cols2_size = [1, 1, 1]
    cols2 = cols[1].columns(cols2_size)
    cols2[0].write('$b$, см')
    b = cols2[1].number_input(label='$b$, см', step=5.0, format="%.1f", value=20.0, min_value=1.0, max_value=500.0, label_visibility="collapsed")
    b = round(b,2)

    cols2 = cols[1].columns(cols2_size)
    cols2[0].write('$h$, см')
    h = cols2[1].number_input(label='$h$, см', step=5.0, format="%.1f", value=40.0, min_value=1.0, max_value=500.0, label_visibility="collapsed")
    h = round(h,2)

    cols2 = cols[1].columns(cols2_size)
    cols2[0].write('$h_0$, см')
    h0 = cols2[1].number_input(label='$h_0$, см', step=5.0, format="%.1f", value=20.0, min_value=1.0, max_value=500.0, label_visibility="collapsed")
    h0 = round(h0,2)

    cols2 = cols[1].columns(cols2_size)
    cols2[0].write('$c_L$, см')
    is_cL = cols2[2].toggle('Контур_слева', value=True, label_visibility="collapsed")
    cL = cols2[1].number_input(label='$c_L$, см', step=5.0, format="%.1f", value=0.0, disabled=is_cL, min_value=0.0, max_value=500.0, label_visibility="collapsed")
    cL = round(cL,2)


    cols2 = cols[1].columns(cols2_size)
    cols2[0].write('$c_R$, см')
    is_cR = cols2[2].toggle('Контур_справа', value=True, label_visibility="collapsed")
    cR = cols2[1].number_input(label='$c_R$, см', step=5.0, format="%.1f", value=0.0, disabled=is_cR, min_value=0.0, max_value=500.0, label_visibility="collapsed")
    cR = round(cR,2)


    cols2 = cols[1].columns(cols2_size)
    cols2[0].write('$c_B$, см')
    is_cB = cols2[2].toggle('Контур_снизу', value=True, label_visibility="collapsed")
    cB = cols2[1].number_input(label='$c_B$, см', step=5.0, format="%.1f", value=0.0, disabled=is_cB, min_value=0.0, max_value=500.0, label_visibility="collapsed")
    cB = round(cB,2)


    cols2 = cols[1].columns(cols2_size)
    cols2[0].write('$c_T$, см')
    is_cT = cols2[2].toggle('Контур_сверху', value=True, label_visibility="collapsed")
    cT = cols2[1].number_input(label='$c_T$, см', step=5.0, format="%.1f", value=0.0, disabled=is_cT, min_value=0.0, max_value=500.0, label_visibility="collapsed")
    cT = round(cT,2)

    cols2 = cols[1].columns(cols2_size)
    cols2[0].write('$q, тс/м^2$')
    is_q = cols2[2].toggle('Отпор', value=False, label_visibility="collapsed")
    if is_q == True: dis_q = False
    else: dis_q = True
    q = cols2[1].number_input(label='$q, тс/м^2$', step=5.0, format="%.2f", value=1.15, disabled=dis_q, min_value=0.0, max_value=500.0, label_visibility="collapsed")
    if is_q == True: q = round(q, 2)
    else: q = 0.0
    

    cols2 = st.columns([1,0.6,0.6,0.7,0.7,0.7,0.5,0.5])
    ctype = cols2[0].selectbox(label='Бетон', options=available_concretes, index=5, label_visibility="visible")
    selected_concrete_data = table_concretes_data.loc[table_concretes_data['Class'] == ctype]
    selected_concrete_data = selected_concrete_data.to_dict('records')[0]
    Rbt0 = selected_concrete_data['Rbt']
    Rbt0 = round(Rbt0,4)
    gammabt = cols2[1].number_input(label='$\\gamma_{bt}$', step=0.05, format="%.2f", value=1.0, min_value=0.1, max_value=1.0, label_visibility="visible")
    gammabt = round(gammabt,2)
    Rbt01 = Rbt0*gammabt
    RbtMPA = cols2[2].number_input(label='$R_{bt}$, МПа', step=0.05, format="%.2f", value=Rbt01, min_value=0.1, max_value=2.2, label_visibility="visible", disabled=True)
    RbtMPA = round(RbtMPA,4)
    Rbt = 0.01019716213*RbtMPA
    #Rbt = 0.01*RbtMPA
    Rbt = round(Rbt,5)
    F0 = cols2[3].number_input(label='$F$, тс', step=1.0, format="%.1f", value=49.0, min_value=1.0, max_value=50000.0, label_visibility="visible")
    Mxloc = cols2[4].number_input(label='$M_{x,loc}$, тсм', step=0.5, format="%.2f", value=4.0, label_visibility="visible")
    Myloc = cols2[5].number_input(label='$M_{y,loc}$, тсм', step=0.5, format="%.2f", value=7.0, label_visibility="visible")
    deltaMx = cols2[6].number_input(label='$\\delta_{Mx}$', step=0.1, format="%.2f", value=0.5, min_value=0.0, max_value=2.0, label_visibility="visible")
    deltaMy = cols2[7].number_input(label='$\\delta_{My}$', step=0.1, format="%.2f", value=0.5, min_value=0.0, max_value=2.0, label_visibility="visible")
    cols2 = st.columns([1.1, 1.1, 1.1, 1, 1])
    delta_M_exc = cols2[-1].selectbox(label='$\\delta_{M}$ для $F \\cdot e$', options=['нет','да'], index=0, label_visibility="visible")
    if delta_M_exc == 'да': delta_M_exc=True
    else: delta_M_exc=False
    M_abs = cols2[-3].selectbox('Учитывать знак $F \\cdot e$', options=['нет', 'да'], index=0)
    if M_abs == 'да': M_abs=False
    else: M_abs=True
    F_dir = cols2[-2].selectbox('Направление $F$', options=['вверх','вниз'], index=0)
    if F_dir == 'вверх':
        image_f_dir = Image.open("F_up.png")
    if F_dir == 'вниз':
        image_f_dir = Image.open("F_down.png")
    st.sidebar.write("Положительные направления нагрузок")
    st.sidebar.image(image_f_dir)
    is_sw = cols2[0].selectbox(label='Поперечное арм.', options=['да', 'нет'], index=0, label_visibility="visible")
    if is_sw == 'да': is_sw=True
    else: is_sw=False
    sw_mode = cols2[1].selectbox(label='Режим арм.', options=['подбор','проверка'], index=0, label_visibility="visible")
    if sw_mode == 'подбор': sw_block=True
    else: sw_block=False


    if is_sw:
        cols2 = st.columns([1,0.8,1,1,1,1,1])
        rtype = cols2[0].selectbox(label='Арматура', options=['A240', 'A400', 'A500'], index=0, label_visibility="visible")
        selected_reinf_data = table_reinf_data.loc[table_reinf_data['Class'] == rtype]
        selected_reinf_data = selected_reinf_data.to_dict('records')[0]
        Rsw0 = selected_reinf_data['Rsw']
        Rsw = cols2[1].number_input(label='$R_{sw}$, МПа', value=Rsw0, label_visibility="visible", disabled=True)
        Rsw = 0.01019716213*Rsw
        Rsw = round(Rsw,3)
        nsw = cols2[3].number_input(label='$n_{sw}$, шт.', step=1, format="%i", value=2, min_value=1, max_value=10, label_visibility="visible")
        nsw = round(nsw)
        sw = cols2[5].number_input(label='$s_w$, см', step=5.0, format="%.2f", value=6.0, min_value=0.1, max_value=100.0, label_visibility="visible", disabled=sw_block)
        sw = round(sw,2)
        dsw = cols2[4].selectbox(label='$d_{sw}$, мм', options=dias, index=0, label_visibility="visible")
        dsw = round(dsw)
        Asw = pi*dsw*dsw/4*nsw/10/10
        Asw= round(Asw,3)
        qsw0 = Rsw*Asw/sw
        qsw0 = round(qsw0,5)
        qsw_max = (Rbt*h0)/0.8
        qsw_max = round(qsw_max,5)
        qsw_min = 0.25*Rbt*h0/0.8
        #qsw_min = round(qsw_min,3)
        #st.write(qsw_max)
        ksw0 = 0.8*qsw0/(Rbt*h0)
        ksw0 = round(ksw0,3)
        cols2[6].number_input(label='$k_{sw}$, %', step=5.0, format="%.1f", value=ksw0*100, label_visibility="visible", disabled=True)
        kh0 = cols2[2].number_input(label='$k \\cdot h_0$', step=0.1, format="%.1f", value=1.5, label_visibility="visible")
        ksw = 0.0
        if ksw0<0.25:
            ksw = 0.0
            qsw = 0
            sw_max = Rsw*pi*dsw*dsw*nsw/4/100/qsw_min
            sw_max = round(sw_max,1)
        if ksw0>1.0:
            ksw = 1.0
            qsw = qsw_max
        if 0.25<=ksw0<=1:
            ksw = ksw0
            qsw = qsw0
        ksw = round(ksw,3)

#num_elements = 2
#Основной расчет (новый)
if True:
    solve_data = {'b': b, 'h': h, 'dh0': h0, 'h0': h0, 'kh0': kh0,
                  'is_cL': is_cL, 'is_cR': is_cR, 'is_cB': is_cB, 'is_cT': is_cT,
                  'cL': cL, 'cR': cR, 'cB': cB, 'cT': cT,
                  'Rbt': Rbt, 'ctype': ctype, 'Rbt0': Rbt0, 'gammabt': gammabt, 'RbtMPA': RbtMPA,
                  'F0': F0, 'Mxloc': Mxloc, 'Myloc': Myloc, 'deltaMx': deltaMx, 'deltaMy': deltaMy,
                  'F_dir': F_dir, 'delta_M_exc': delta_M_exc, 'M_abs': M_abs, 'xF': round(b/2,2), 'yF': round(h/2,2), 'q': q,
                  'is_sw': is_sw, 'Rsw': Rsw, 'Rsw0': Rsw0, 'Asw': Asw, 'sw': sw, 'sw_mode': sw_mode, 'rtype': rtype, 'dsw': dsw, 'nsw': nsw, 'qsw0': qsw0,
                  'is_geom_full': is_geom_full}
    #Результаты расчета по первому контуру
    result = single_solve(**solve_data)
    #Число элементов в контуре
    num_elements = len(result['contours'])
    #Результаты расчета по второму контуру
    solve_data_second = solve_data.copy()
    solve_data_second.update({'dh0': 2*kh0*h0})
    result_second = single_solve(**solve_data_second)
    

fig, image_stream = draw_geometry(result) #геометрия для расчета по первому контуру
fig_sw, image_stream_sw = draw_geometry(result_second) #геометрия для расчета за поперечной арматурой

fig_Aq, image_stream_Aq = None, None
if result['Fq'] != 0.0:
    fig_Aq, image_stream_Aq = draw_Aq(result)
    st.sidebar.write('Габариты основания пирамиды')
    st.sidebar.plotly_chart(fig_Aq)

cols[0].plotly_chart(fig)

#st.plotly_chart(fig_sw)

if num_elements<2:
    st.subheader('В расчете должно быть минимум два участка!')
    st.stop()

#st.write(result)
#Быстрый вывод основных результатов
string = 'По бетону '
if result['kb']>2: string += ':red[$k_b = ' + str(result['kb']) + '$]'
if result['kb']<=1: string += ':green[$k_b = ' + str(result['kb']) + '$]'
if 1<result['kb']<=2: string += ':orange[$k_b = ' + str(result['kb']) + '$]'
if result['is_sw']:
    if result['sw_mode'] == 'проверка' or (result['sw_mode'] == 'подбор' and 2 >= result['kb'] > 1):
        if result['k']>1  and result['kb'] > 2: string += ', суммарно :red[$k=' + str(result['k']) + '$]'
        if result['k']<=1: string += ', суммарно :green[$k=' + str(result['k']) + '$]'
        if result['k']>1 and result['kb']<=2: string += ', суммарно :orange[$k=' + str(result['k']) + '$]'
        string += ', за поперечным армированием '
        if result_second['kb']>1: string += ':red[$k=' + str(result_second['kb']) + '$]'
        if result_second['kb']<1: string += ':green[$k=' + str(result_second['kb']) + '$]'
string += '.'
st.write(string)

with st.expander('Краткий отчет'):
    fast_report (result, result_second)

if is_report:
    doc = Document('Template.docx')
    styles_element = doc.styles.element
    rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
    lang_default = rpr_default.xpath('w:lang')[0]
    lang_default.set(docx.oxml.shared.qn('w:val'),'ru-RU')
    doc.core_properties.author = 'Автор'
    p = doc.paragraphs[-1]
    delete_paragraph(p)
    string = 'Расчет на продавливание при действии сосредоточенной силы и изгибающих моментов.'
    doc.add_heading(string, level=1)
    #st.subheader(string)

    string = 'Расчет производится согласно СП 63.13330.2018 п. 8.1.46 – 8.1.52.'
    if is_sw: string += ' Поперечная арматура принимается равномерно расположенной по периметру расчетного контура.'
    doc.add_paragraph().add_run(string)
    string = 'Геометрические характеристики, такие как осевые моменты инерции и моменты сопротивления для расчетного контура вычисляются в направлении соответствующих осей.'
    string += ' Принятые в расчете единицы измерения: длина – $см$; сила – $тс$; напряжение – $тс/см^2$.'
    #st.write(string)
    add_text_latex(doc.add_paragraph(), string)


    if is_init_help: report_solve_method(result, doc) #Вывод справки по расчету при необходимости
    report_init_data(result, doc) #Вывод исходных данных

   
    #Расчет геометрических характеристик
    if result['is_geom_full']: report_full_geometry(result, doc, image_stream, image_stream_Aq)
    else: report_short_geometry(result, doc, image_stream, image_stream_Aq)

    #Расчет действующих усилий
    report_acting_forces(result, doc)
    
    #Расчет предельных усилий, воспринимаемых бетоном
    report_concrete_ultimate_forces(result, doc)

    if result['is_sw'] and result['sw_mode'] == 'проверка':
        #Расчет предельных усилий, воспринимаемых арматурой
        report_reinf_ultimate_forces(result, doc)
        report_full_strength(result, doc)
    
    #Проверка прочности по бетону
    if result['sw_mode'] == 'подбор':
        report_concrete_strength(result, doc)

    #Вывод блока подбора шага поперечного армирования при необходимости
    if result['is_sw'] and result['sw_mode'] == 'подбор' and 2>= result['kb'] > 1:
        #Подбор шага поперечного армирования
        report_solve_sw_min(result, doc)
        #Расчет предельных усилий, воспринимаемых арматурой
        report_reinf_ultimate_forces(result, doc)
        report_full_strength(result, doc)

   
    #st.plotly_chart(fig_sw)




    file_stream = io.BytesIO()
    doc.save(file_stream)
    st.sidebar.download_button('Отчет *.docx', file_name='продавливание_колонна_' + str(round(b))+'x'+str(round(h))+'_h0='+str(round(h0)) +'.docx', data=file_stream)







