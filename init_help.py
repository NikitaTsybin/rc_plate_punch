import streamlit as st
import typing
import latex2mathml.converter
import mathml2omml
from docx.oxml import parse_xml
from docx.shared import Mm
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from word_func import *



def init_data():
    with st.expander('Пояснения для исходных данных'):
        st.write('$b$ и $h$ – ширина и высота поперечного сечения сечения колонны, см;')
        st.write('$h_0$ – приведенная рабочая высота поперечного сечения плиты, см;')
        st.write('$c_L$ и $c_R$ – расстояние в свету до левой и правой грани плиты от грани колонны, см. Вводится если левая или правая граница контура отключена;')
        st.write('$c_B$ и $c_T$ – расстояние в свету до нижней и верхней грани плиты от грани колонны, см. Вводится если верхняя или нижняя граница контура отключена;')
        st.write('$R_{bt}$ – расчетное сопротивление на растяжение бетона, МПа;')
        st.write('''$\\gamma_{bt}$ – коэффициент, вводимый к расчетному сопротивлению бетона на растяжение.
        Например, в соответствии с п. 6.1.12 (а) СП 63.13330.2018 при продолжительном действии нагрузки величина данного коэффициента принимается равной 0.9;''')
        st.write('$R_{sw}$ – расчетное сопротивление поперечного армирования, МПа;')
        st.write('$n_{sw}$ – число стержней поперечного армирования в одном ряду, пересекающих пирамиду продавливания, шт.;')
        st.write('$s_{w}$ – шаг рядов поперечного армирования вдоль расчетного контура, мм;')
        st.write('$d_{sw}$ – диаметр поперечного армирования, мм;')
        st.write('$k_{sw}, %$ – вклад поперечного армирования в несущую способность в процентах от максимально допустимого;')
        st.write('''$F$ – сосредоточенная продавливающая сила, тс.
        Значение сосредоточенной силы следует принимать за вычетом сил, действующих в пределах основания
        пирамиды продавливания в противоположном направлении; ''')
        st.write('$M_{x,loc}$ – сосредоточенный момент в месте приложения сосредоточенной нагрузки в ПЛОСКОСТИ оси $x$ (относительно оси $y$), тсм. Положительное значение "сжимает" правую грань расчетного контура (см. "положительные направления нагрузок");')
        st.write('$M_{y,loc}$ – сосредоточенный момент в месте приложения сосредоточенной нагрузки в ПЛОСКОСТИ оси $y$ (относительно оси $x$), тсм. Положительное значение "сжимает" верхнюю грань расчетного контура (см. "положительные направления нагрузок");')
        st.write('$\\delta_{Mx}$ – понижающий коэффициент к моментам в плоскости оси $x$;')
        st.write('$\\delta_{My}$ – понижающий коэффициент к моментам в плоскости оси $y$;')
        st.write('''Параметр "учитывать знак $F \\cdot e$".
        Если данный параметр включен, то момент от внецентренного приложения сосредоточенной силы,
        в зависимости от величины эксцентриситета, может догружать или разгружать расчетный контур.
        В противном случае момент от эксцентриситета всегда догружает расчетный контур;''')
        st.write('''Параметр "направление $F$". Данный параметр влияет на результаты расчета, только если активирован параметр "учитывать знак $F \\cdot e$". При положительном
        значении эксцентриситета и силе направленной вверх, момент от эксцентриситета $F \\cdot e$ будет положительным.
        Если сила направлена вниз – отрицательным. Более наглядно этот факт понятен из рисунка "положительные направления нагрузок";''')
        st.write('''Параметр "$\\delta_{M}$ для $F \\cdot e$". Данный параметр указывает, учитывать ли понижающие коэффициенты $\\delta_{M}$ для моментов от эксцентриситета.
        В соответствии с п. 8.1.46 указания по снижению момента представлены только в абзаце, содержащем $M_{loc}$.''')


def init_help(doc, is_sw, M_abs, delta_M_exc, F_dir):
    with st.expander('Расчетные предпосылки'):
        string = 'Расчетные предпосылки.'
        doc.add_heading(string, level=1)
        st.subheader(string)

        string = 'Расчет производится согласно СП 63.13330.2018 п. 8.1.46 – 8.1.52.'
        if is_sw: string += ' Поперечная арматура принимается равномерно расположенной по периметру расчетного контура.'
        st.write(string)
        
        string = 'Проверка прочности в общем случае выполняется из условия (8.96):'
        st.write(string)
        doc.add_paragraph().add_run(string)
        string = '$\\dfrac{F}{F_{b,ult}+F_{sw,ult}} + \\dfrac{M_x}{M_{bx,ult} + M_{sw,x,ult}} + \\dfrac{M_y}{M_{by,ult}+ M_{sw,y,ult} }  \\le 1.0.$'
        st.write(string)
        add_math(doc.add_paragraph(), string.replace('$',''))

        string = 'Здесь $F$, $M_{x}$ и $M_{y}$ – сосредоточенная сила и изгибающие моменты от внешней нагрузки, учитываемые при расчете на продавливание, определяемые в соответствии с п. 8.1.46;'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$F_{b,ult}$, $M_{bx,ult}$ и $M_{by,ult}$ – предельные сосредоточенная сила и изгибающие моменты, которые могут быть восприняты бетоном в расчетном поперечном сечении;'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = '$F_{sw,ult}$, $M_{sw,x,ult}$ и $M_{sw,y,ult}$ – предельные сосредоточенная сила и изгибающие моменты, которые могут быть восприняты арматурой в расчетном поперечном сечении.'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = 'Сосредоточенные моменты от внешней нагрузки $M_{x}$ и $M_{y}$, учитываемые при расчете на продавливание, определяются в соответствии указаниями с п. 8.1.46.'
        string += ' В общем случае, при расположении сосредоточенной силы $F$ внецентренно относительно центра тяжести контура расчетного поперечного сечения'
        string += ' с эксцентриситетами $e_x$ и $e_y$ по формуле:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        if M_abs:
            if delta_M_exc:
                string = '$M_x = (|M_{x,loc}| +  F \\cdot | e_x |) \\cdot \\delta_{Mx}; \\quad'
                string += ' M_y = (|M_{y,loc}| +  F \\cdot | e_y |) \\cdot \\delta_{My}.$'
            else:
                string = '$M_x = |M_{x,loc}| \\cdot \\delta_{Mx} +  F \\cdot | e_x| ; \\quad'
                string += ' M_y =|M_{y,loc}| \\cdot \\delta_{My} +  F \\cdot | e_y|.$'
        else:
            if F_dir == 'вверх': znak = '+'
            else: znak = '-'
            if delta_M_exc:
                string = '$M_x = |M_{x,loc}' + znak + 'F \\cdot e_x| \\cdot \\delta_{Mx}; \\quad'
                string += ' M_y = |M_{y,loc} ' + znak + ' F \\cdot e_y| \\cdot \\delta_{My}.$'
            else:
                string = '$M_x = |M_{x,loc} \\cdot \\delta_{Mx} ' + znak + ' F \\cdot e_x| ; \\quad'
                string += ' M_y = |M_{y,loc} \\cdot \\delta_{My} ' + znak + ' F \\cdot e_y| .$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'В данных формулах $M_{loc}$ – действующий момент в месте приложения сосредоточенной нагрузки.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Коэффициенты $\\delta_{Mx} \\le 1$ и $\\delta_{My} \\le 1$ учитывают положения п. 8.1.46:'
        string += ' при действии момента в месте приложения сосредоточенной нагрузки половину этого момента учитывают при расчете на продавливание, а другую половину – при расчете по нормальным сечениям.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Предельная продавливающая сила, воспринимаемая бетоном $F_{b,ult}$, вычисляется по формуле (8.88) с учетом формулы (8.89):'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = '$F_{b,ult} = R_{bt} \\cdot h_0 \\cdot u.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Здесь $u$ и $h_0$ – периметр контура расчетного поперечного сечения и приведенная рабочая высота сечения, вычисляемая по формуле:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$h_0 = 0.5 \\cdot (h_{0x} + h_{0y}),$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'где $h_{0x}$ и $h_{0y}$ – рабочая высота сечения для продольной арматуры, расположенной в направлении осей $x$ и $y$.'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = 'Предельная продавливающая сила, воспринимаемая арматурой $F_{sw,ult}$, при равномерном расположении арматуры вдоль контура расчетного поперечного сечения вычисляется по формуле (8.91):'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = '$F_{sw,ult} = 0.8 \\cdot q_{sw} \\cdot u,$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'где $q_{sw}$ – усилие в поперечной арматуре на единицу длины контура расчетного поперечного сечения,'
        string += ' расположенной в пределах расстояния $0.5 \\cdot h_0$ по обе стороны от контура расчетного сечения'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = '$q_{sw} = \\dfrac{R_{sw} \\cdot A_{sw} }{s_w};$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$A_{sw}$ – площадь сечения всей поперечной арматуры расположенная в пределах расстояния $0.5 \\cdot h_0$ по обе'
        string += ' стороны от контура расчетного поперечного сечения по периметру контура расчетного поперечного сечения (площадь поперечной арматуры в одном каркасе);'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = '$s_w$ – шаг поперечной арматуры (шаг каркасов);'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = '$u$ – периметр контура расчетного поперечного сечения.'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = 'Предельный изгибающий момент, воспринимаемый бетоном $M_{b,ult}$, вычисляется по формуле (8.94):'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = '$M_{b,ult} = R_{bt} \\cdot W_b \\cdot h_0.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = 'где $W_b$ – момент сопротивления расчетного поперечного сечения, определяемый в соответствии с п. 8.1.51.'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = 'Усилия $M_{sw,x,ult}$ и $M_{sw,y,ult}$, воспринимаемые поперечной арматурой, нормальной к продольной оси элемента и'
        string += ' расположенной равномерно вдоль контура расчетного сечения, определяются при действии изгибающего момента'
        string += ' соответственно в направлении осей $x$ и $y$ по формуле (8.97):'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$M_{sw,ult} = 0.8 \\cdot q_{sw} \\cdot W_{sw}.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'В дальнейших расчетах, в соответствии с п. 8.1.52, предполагаем, что'
        string += ' поперечная арматура расположена равномерно вдоль расчетного контура продавливания в пределах зоны, границы'
        string += ' которой отстоят на расстоянии $h_0/2$ в каждую сторону от контура продавливания бетона.'
        string += ' В этом случае значения моментов сопротивления поперечной арматуры при продавливании $W_{sw,x(y)}$ принимают равными соответствующим значениям $W_{bx}$ и $W_{by}$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Введем следующей коэффициент:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$k_{sw} = \\dfrac{F_{sw,ult}}{F_{b,ult}} = \\dfrac{M_{sw,ult}}{M_{b,ult}} = \\dfrac{0.8 \\cdot q_{sw}}{R_{bt} \\cdot h_0}.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Данный коэффициент можно интерпретировать, как вклад поперечного армирования в несущую способность. C учетом данного коэффициента и предыдущих выкладок условие прочности (8.96) можно переписать в виде:'
        st.write(string)
        doc.add_paragraph().add_run(string)
        string = '$\\dfrac{F}{F_{b,ult} \\cdot (1+k_{sw})} + \\dfrac{M_x}{M_{bx,ult} \\cdot (1+k_{sw})} + \\dfrac{M_y}{M_{by,ult} \\cdot (1+k_{sw}) }  \\le 1.0.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
    
        string = 'В случае, если поперечная арматура не учитывается в расчете, $k_{sw}=0.0$.'
        string += ' В противном случае, если поперечная арматура учитывается в расчете, на коэффициент $k_{sw}$ накладываются следующие ограничения:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$0.25 \\le k_{sw} \\le 1.0.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Нижняя и верхняя граница $k_{sw}$ связаны со ограничениями, приведенными в п. 8.1.48 и 8.1.50.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Во-первых, поперечную арматуру можно учитывать в расчете при выполнении условия $F_{sw,ult} \\ge 0.25 \\cdot F_{b,ult}$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Во-вторых, усилия, воспринимаемые поперечной арматурой, не могут превышать усилия, воспринимаемые бетоном, т.е.'
        string += ' $F_{b,ult} + F_{sw,ult} \\le 2 \\cdot F_{b,ult}$,'
        string += ' $M_{bx,ult} + M_{sw,x,ult} \\le 2 \\cdot M_{bx,ult}$,'
        string += ' $M_{by,ult} + M_{sw,y,ult} \\le 2 \\cdot M_{by,ult}$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Введем дополнительно коэффициенты $k_{b,F}$, $k_{b,M}$ и $k_b$ следующим образом:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$k_{b,F} = \\dfrac{F}{F_{b,ult}}; \\quad'
        string += ' k_{b,M} = \\dfrac{M_x}{M_{bx,ult}} + \\dfrac{M_y}{M_{by,ult}}; \\quad k_b = k_{b,F} + k_{b,M}$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'В результате, условие прочности примет вид:'
        st.write(string)
        doc.add_paragraph().add_run(string)
        string = '$\\dfrac{k_{b,F} + k_{b,M}}{1+k_{sw}} \\le 1.0.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Здесь $k_{b,F}$, $k_{b,M}$ и $k_b$ – коэффициенты использования прочности бетона расчетного поперечного сечения по силе, моментам и суммарный соответственно.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'В соответствии с п. 8.1.46 при действии сосредоточенных моментов и силы в условиях прочности соотношение между действующими'
        string += ' сосредоточенными моментами $M$, учитываемыми при продавливании, и предельными $M_{ult}$ принимают не более'
        string += ' половины соотношения между действующим сосредоточенным усилием $F$ и предельным $F_{ult}$.'
        string += ' Следовательно, $k_{b,M} \\le 0.5 \\cdot k_{b,F}$. В случае нарушения данного условия значение $k_{b,M}$ принимается равным $k_{b,M}=0.5 \\cdot k_{b,F}$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Условие прочности, в случае $k_{b,M} \\ge 0.5 \\cdot k_{b,F}$ примет вид:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$\\dfrac{1.5 \\cdot k_{b,F}}{1+k_{sw}} \\le 1.0.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Из данного неравенства можно получить условие прочности, справедливое для любого значения моментов:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$k_{b,F} \\le \\dfrac{1+k_{sw}}{1.5} \\Rightarrow F \\le \\dfrac{F_{b,ult} \\cdot (1+k_{sw})}{1.5}.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Аналогично при отсутствии поперечной арматуры ($k_{sw}=0.0$):'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$k_{b,F} \\le \\dfrac{1}{1.5} \\Rightarrow F \\le \\dfrac{F_{b,ult}}{1.5}.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Для расчета используется следующая последовательность.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '1. Выполняется расчет геометрических характеристик контура расчетного поперечного сечения.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '2. Выполняется расчет предельных усилий, воспринимаемых бетоном расчетного поперечного сечения.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '3. Вычисление коэффициентов $k_{b,F}$ и $k_{b,M}$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '4. Проверка условия $k_{b,M} \\le 0.5 \\cdot k_{b,F}$. Если условие не выполняется, коэффициент $k_{b,M}$ принимается равным $k_{b,M} = 0.5 \\cdot k_{b,F}$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '5. Вычисление коэффициента использования прочности бетона расчетного поперечного сечения $ k_{b} = k_{b,F} + k_{b,M}$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '6. В случае, если $k_{b} \\le 1.0 $ прочность обеспечена без установки поперечной арматуры.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '7. В случае, если $k_{b} > 2.0 $ прочность не может быть обеспечена, необходимо увеличение габаритов площадки передачи нагрузки, либо толщины плиты, либо класса бетона.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '8. В случае, если $1.0 < k_{b} \\le 2.0 $ требуется установка поперечной арматуры.'
        string += ' В этом случае требуемое соотношение $A_{sw}/s_w$, с учетом выражений для $k_{sw}$ и $q_{sw}$ определяется по формуле:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$\\dfrac{A_{sw}}{s_w} \\ge \\dfrac{R_{bt} \\cdot h_0 \\cdot (k_{b} - 1)}{0.8 \\cdot R_{sw}}.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
