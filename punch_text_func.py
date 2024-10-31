import streamlit as st
import pandas as pd
from docx.shared import Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from punch_word_func import *
import os
import sys


add_dir1 = ''
##add_dir1 = (os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
##+ '/punch/')
##sys.path.append(add_dir)

def fast_report (result, result_second): #Быстрый отчет
    st.write('Расчетные усилия:')
    string = '$F = ' + str(result["F"]) +  '\\cdot тс; '
    string += '\\quad M_x = ' + str(result["Mx"])+  ' тсм; '
    string += '\\quad M_y = ' + str(result["My"])+  ' тсм.$'
    if result['Fq'] != 0.0:
        string += ' :blue[Учтена разгружающая сила $F_q = q \\cdot A_q = ' + str(result['q']) + '\\cdot' + str(result['Aq']) + '=' + str(result['Fq']) + ' тс$.]'
    st.write(string)
    st.write('Предельные усилия, воспринимаемые бетоном:')
    string = '$F_{b,ult} = ' + str(result["Fbult"]) +  '\\cdot тс; '
    string += '\\quad M_{bx,ult} = ' + str(result["Mbxult"])+  ' тсм; '
    string += '\\quad M_{by,ult} = ' + str(result["Mbyult"])+  ' тсм.$'
    st.write(string)
    st.write('Коэффициенты использования бетона по силе $k_{b,F}$, по моментам $k_{b,M}$ и суммарный $k_{b}$:')
    string = '$k_{b,F}='  + str(result['kbF'])
    string += '; \\quad k_{b,M}='  + str(result['kbM'])
    string += '; \\quad k_{b}='  + str(result['kb']) + '$.'
    if result['kbM0'] != result['kbM']:
        string += ' :blue[Вклад моментов ограничен.]'
    if result['kb'] > 2:
        string += ' :red[Прочность не может быть обеспечена, необходимо увеличение габаритов колонны, толщины плиты или класса бетона.]'
    if result['kb'] <= 1:
        string += ' :green[Прочность обеспечена. Поперечное армирование не требуется.]'
    if 2>= result['kb'] > 1:
        string += ' :orange[Требуется поперечное армирование.]'
    st.write(string)
    if result['is_sw']:
        if 2>= result['kb']:
            if result['sw_mode'] == 'подбор' and result['kb']>1:
                string = 'Максимальный шаг, при заданном $A_{sw} = ' + str(result['Asw']) + ' см^2$, составляет $s_w = ' + str(result['sw']) + ' см$.'
                if result['sw_min_code'] == 1:
                    string += ' :blue[Учтено ограничение на максимальный шаг.]'
                if result['kb_sw_code'] == 1:
                    string += ' :blue[Учтено требование $F_{sw,ult} \\ge 0.25 \\cdot F_{b,ult}$.]'
                string += ' При данном шаге $q_{sw} = ' + str(result['qsw']) + ' тс/см$.'
                string += ' Вклад поперечного армирования ' + str(round(result['ksw']*100,1)) + '% от максимального.'
                st.write(string)
            if result['sw_mode'] == 'проверка':
                string = 'При заданном $A_{sw} = ' + str(result['Asw']) + ' см^2$ и шаге $s_w = ' + str(result['sw']) + ' см$ '
                string += 'усилие в поперечной арматуре $q_{sw} = ' + str(result['qsw']) + ' тс/см$.'
                if result['ksw']>1:
                    string += ':orange['
                string += ' Вклад поперечного армирования ' + str(round(result['ksw']*100,1)) + '% от максимального.'
                if result['ksw']>1:
                    string += ']'
                st.write(string)
            if result['sw_mode'] == 'проверка' or (result['sw_mode'] == 'подбор' and result['kb']>1):
                st.write('Предельные усилия, воспринимаемые арматурой:')
                string = '$F_{sw,ult} = ' + str(result["Fswult"]) +  '\\cdot тс; '
                string += '\\quad M_{sw,x,ult} = ' + str(result["Mswxult"])+  ' тсм; '
                string += '\\quad M_{sw,y,ult} = ' + str(result["Mswyult"])+  ' тсм.$'
                st.write(string)
                if result['Fsw_code'] == 1:
                    string = ':blue[Вклад поперечного армирования ограничен несущей способностью бетона]'
                    st.write(string)
                st.write('Коэффициенты использования по силе $k_F$, по моментам $k_M$ и суммарный $k$:')
                string = '$k_{F}='  + str(result['kF'])
                string += '; \\quad k_{M}='  + str(result['kM'])
                string += '; \\quad k='  + str(result['k']) + '$.'
                if result['kM0'] != result['kM']:
                    string += ' :blue[Вклад моментов ограничен.]'
                st.write(string)
                if result['k'] < 1:
                    string = ':green[Прочность обеспечена.]'
                if result['k'] > 1:
                    string = ':red[Прочность не обеспечена.]'
                st.write(string)
        if result['sw_mode'] == 'проверка' or (result['sw_mode'] == 'подбор' and result['kb']>1):
            string = 'Проверка за зоной установки поперечной арматуры.'
            string += ' Предельные усилия, воспринимаемые бетоном:'
            st.write(string)
            string = '$F_{b,ult} = ' + str(result_second["Fbult"]) +  '\\cdot тс; '
            string += '\\quad M_{bx,ult} = ' + str(result_second["Mbxult"])+  ' тсм; '
            string += '\\quad M_{by,ult} = ' + str(result_second["Mbyult"])+  ' тсм.$'
            st.write(string)
            string = ' Коэффициенты для расчетного контура на расстоянии $' + str(result_second['kh0']) + '\\cdot h_0=' + str(round(result_second['h0']*result_second['kh0'],1)) +  ' см$:'
            st.write(string)
            string = '$k_{b,F}='  + str(result_second['kbF'])
            string += '; \\quad k_{b,M}='  + str(result_second['kbM'])
            string += '; \\quad k_{b}='  + str(result_second['kb']) + '$.'
            if result_second['kbM0'] != result_second['kbM']:
                string += ' :blue[Вклад моментов ограничен.]'
            if result['Fq'] != 0.0:
                string += ' :blue[Учтена разгружающая сила $F_q = q \\cdot A_q = ' + str(result_second['q']) + '\\cdot' + str(result_second['Aq']) + '=' + str(result_second['Fq']) + ' тс$.]'
            st.write(string)
            if result_second['kb'] <= 1:
                string = ':green[Прочность за зоной поперечного армирования обеспечена.]'
                st.write(string)
            if result_second['kb'] > 1:
                string = ':orange[Требуется увеличение зоны поперечного армирования.]'
                st.write(string)
            #st.write(result_second['kbF'], result_second['kbM'], result_second['kb'])

def init_data_help(): #Пояснения к параметрам
    with st.expander('Пояснения для исходных данных'):
        st.write('$b$ и $h$ – ширина и высота поперечного сечения сечения колонны, см;')
        st.write('$h_0$ – приведенная рабочая высота поперечного сечения плиты, см;')
        st.write('$c_L$ и $c_R$ – расстояние в свету до левой и правой грани плиты от грани колонны, см. Вводится если левая или правая граница контура отключена;')
        st.write('$c_B$ и $c_T$ – расстояние в свету до нижней и верхней грани плиты от грани колонны, см. Вводится если верхняя или нижняя граница контура отключена;')
        st.write('$R_{bt}$ – расчетное сопротивление на растяжение бетона, МПа;')
        st.write('''$\\gamma_{bt}$ – коэффициент, вводимый к расчетному сопротивлению бетона на растяжение.
        Например, в соответствии с п. 6.1.12 (а) СП 63.13330.2018 при продолжительном действии нагрузки величина данного коэффициента принимается равной 0.9;''')
        st.write('$R_{sw}$ – расчетное сопротивление поперечного армирования, МПа;')
        st.write('$k \\cdot h_0$ – расстояние от грани колонны до расчетного контура, проверяемого без учета поперечного армирования в долях от $h_0$ (не менее $1.5 \\cdot h_0$);')
        st.write('$n_{sw}$ – число стержней поперечного армирования в одном ряду, пересекающих пирамиду продавливания, шт.;')
        st.write('$d_{sw}$ – диаметр поперечного армирования, мм;')
        st.write('$s_{w}$ – шаг рядов поперечного армирования вдоль расчетного контура, мм. В режиме подбора данная величина вычисляется при заданных $n_{sw}$ и $d_{sw}$;')
        st.write('$k_{sw}, %$ – вклад поперечного армирования при заданных $n_{sw}$, $d_{sw}$ и $s_w$ в несущую способность в процентах от максимально допустимого;')
        st.write('''$F$ – сосредоточенная продавливающая сила, тс.''')
        st.write('''$q$ – распределенные силы, действующих в пределах основания пирамиды продавливания в противоположном для $F$ направлении, $тс/м^2$;''')
        st.write('$M_{x,loc}$ – сосредоточенный момент в месте приложения сосредоточенной нагрузки в ПЛОСКОСТИ оси $x$ (относительно оси $y$), тсм. Положительное значение "сжимает" правую грань расчетного контура (см. "положительные направления нагрузок");')
        st.write('$M_{y,loc}$ – сосредоточенный момент в месте приложения сосредоточенной нагрузки в ПЛОСКОСТИ оси $y$ (относительно оси $x$), тсм. Положительное значение "сжимает" верхнюю грань расчетного контура (см. "положительные направления нагрузок");')
        st.write('$\\delta_{Mx}$ – понижающий коэффициент к моментам в плоскости оси $x$;')
        st.write('$\\delta_{My}$ – понижающий коэффициент к моментам в плоскости оси $y$;')
        st.write('''Параметр "учитывать знак $F \\cdot e$".
        Если данный параметр включен, то момент от внецентренного приложения сосредоточенной силы,
        в зависимости от величины эксцентриситета, может догружать или разгружать расчетный контур.
        В противном случае момент от эксцентриситета всегда догружает расчетный контур;''')
        st.write('''Параметр "направление $F$". Данный параметр влияет на результаты расчета, только если активирован параметр "учитывать знак $F \\cdot e$". При положительном
        значении эксцентриситета и силе направленной вверх, момент от эксцентриситета $F \\cdot e$ будет положительным.
        Если сила направлена вниз – отрицательным. Более наглядно этот факт понятен из рисунка "положительные направления нагрузок";''')
        st.write('''Параметр "$\\delta_{M}$ для $F \\cdot e$". Данный параметр указывает, учитывать ли понижающие коэффициенты $\\delta_{M}$ для моментов от эксцентриситета.
        В соответствии с п. 8.1.46 указания по снижению момента представлены только в абзаце, содержащем $M_{loc}$.''')

def report_init_data(result, doc): #Печать исходных данных
    with st.expander('Исходные данные'):
        string = 'Исходные данные.'
        doc.add_heading(string, level=2)
        st.subheader(string)

        string = 'Геометрия.'
        #doc.add_heading(string, level=3)
        st.subheader(string)

        string = 'Ширина колонны $b=' + str(result['b']) + ' см$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = 'Высота колонны $h=' + str(result['h']) + ' см$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = 'Приведенная рабочая высота сечения плиты $h_0=' + str(result['h0']) + ' см$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Приложенные усилия.'
        #doc.add_heading(string, level=3)
        st.subheader(string)

        if result['F_dir'] == 'вверх':
            image_f_dir = add_dir1 + "F_up.png"
        if result['F_dir'] == 'вниз':
            image_f_dir = add_dir1 + "F_down.png"

        doc.add_picture(image_f_dir, height=Mm(40))
        p = doc.paragraphs[-1]
        p.paragraph_format.first_line_indent = Mm(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p00 = doc.add_paragraph()
        p00.style = 'Рис. 3'
        p00.add_run('Положительные направления действующих усилий')

        string = 'Сосредоточенная сила $F=' + str(result['F0']) + ' тс$, направленная ' + result['F_dir'] + '.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        if round(abs(result['q']), 1) != 0.0:
            string = 'Распределенные силы $q=' + str(result['q']) +  ' тс/м^2$, действующие в пределах основания пирамиды продавливания в противоположном для $F$ направлении.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
        string = 'Изгибающий момент в плоскости оси $x$ в месте приложения силы $M_{x,loc}=' + str(result['Mxloc']) + ' тсм$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = 'Понижающий коэффициент для'
        if result['delta_M_exc']:
            string += ' всех моментов в плоскости оси $x$, '
        else: string += ' момента $M_{x,loc}$ в плоскости оси $x$, '
        string += '$\\delta_{Mx}=' + str(result['deltaMx']) + '$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Изгибающий момент в плоскости оси $y$ в месте приложения силы $M_{y,loc}=' + str(result['Myloc']) + ' тсм$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = 'Понижающий коэффициент для'
        if result['delta_M_exc']:
            string += ' всех моментов в плоскости оси $y$, '
        else: string += ' момента $M_{y,loc}$ в плоскости оси $y$, '
        string += '$\\delta_{My}=' + str(result['deltaMy']) + '$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)


        string = 'Материалы.'
        #doc.add_heading(string, level=3)
        st.subheader(string)
        string = 'Класс бетона по прочности на сжатие ' + result['ctype'] + ', '
        string += '$R_{bt}=' + str(result['Rbt0']) + ' МПа$, '
        string += '$\\gamma_{bt}=' + str(result['gammabt']) + '$.'
        string += ' Расчетное сопротивление бетона на растяжение, учитываемое в расчете:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string =  '$R_{bt}=R_{bt} \\cdot \\gamma_{bt} = ' + str(result['Rbt0']) + '\\cdot' + str(result['gammabt'])
        string += '=' + str(result['RbtMPA']) + ' МПа='
        string += str(result['Rbt']) + ' тс/см^2'
        string += '.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        if result['is_sw']:
            string = 'Класс поперечной арматуры ' + str(result['rtype']) + '. '
            string += 'Расчетное сопротивление поперечной арматуры, учитываемое в расчете:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string =  '$R_{sw}= ' + str(result['Rsw0']) +  ' МПа='
            string += str(result['Rsw']) + ' тс/см^2'
            string += '.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = 'Поперечное армирование.'
            #doc.add_heading(string, level=3)
            st.subheader(string)
            string = 'Диаметр стержней поперечного армирования ' + '$d_{sw}=' + str(result['dsw']) + ' мм$.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = 'Число стержней в одном ряду, пересекающих пирамиду продавливания ' + '$n_{sw}='  + str(result['nsw']) + ' шт$.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = 'Площадь стержней поперечного армирования одного ряда, пересекающих пирамиду продавливания:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$A_{sw} = \\dfrac{\\pi \\cdot d_{sw}^2 \\cdot n_{sw}}{4} = '
            string += '\\dfrac{\\pi \\cdot ' + str(result['dsw']) + '^2 \\cdot' + str(result['nsw']) + '}{4 \\cdot 100} =' 
            string += str(result['Asw']) + ' см^2.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = 'Примечание. Деление на 100  для перевода $мм^2$ в $см^2$.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

            if str(result['sw_mode']) == 'проверка':
                string = 'Шаг рядов поперечного армирования вдоль расчетного контура ' + '$s_{w}='  +  str(result['sw']) + ' см$.'
                st.write(string)
                add_text_latex(doc.add_paragraph(), string)
                string = 'Усилие в поперечной арматуре на единицу длины контура расчетного поперечного сечения:'
                st.write(string)
                add_text_latex(doc.add_paragraph(), string)
                string = '$q_{sw} = \\dfrac{R_{sw} \\cdot A_{sw} }{s_w} = \\dfrac{' + str(result['Rsw']) + '\\cdot' + str(result['Asw'])
                string += '}' + '{' +  str(result['sw']) +  '}=' + str(result['qsw0']) + ' тс/см.$'
                st.write(string)
                add_text_latex(doc.add_paragraph(), string)

def report_solve_method(result, doc): #Общая теория расчета
    with st.expander('Расчетные предпосылки'):
        string = 'Расчетные предпосылки.'
        doc.add_heading(string, level=2)
        st.subheader(string)

        string = 'Расчет производится согласно СП 63.13330.2018 п. 8.1.46 – 8.1.52.'
        if result['is_sw']: string += ' Поперечная арматура принимается равномерно расположенной по периметру расчетного контура.'
        st.write(string)
        
        string = 'Проверка прочности в общем случае выполняется из условия (8.96):'
        st.write(string)
        doc.add_paragraph().add_run(string)
        string = '$\\dfrac{F}{F_{b,ult}+F_{sw,ult}} + \\dfrac{M_x}{M_{bx,ult} + M_{sw,x,ult}} + \\dfrac{M_y}{M_{by,ult}+ M_{sw,y,ult} }  \\le 1.0.$'
        st.write(string)
        add_math(doc.add_paragraph(), string.replace('$',''))

        string = 'Здесь $F$, $M_{x}$ и $M_{y}$ – сосредоточенная сила и изгибающие моменты от внешней нагрузки, учитываемые при расчете на продавливание, определяемые в соответствии с п. 8.1.46;'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$F_{b,ult}$, $M_{bx,ult}$ и $M_{by,ult}$ – предельные сосредоточенная сила и изгибающие моменты, которые могут быть восприняты бетоном в расчетном поперечном сечении;'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = '$F_{sw,ult}$, $M_{sw,x,ult}$ и $M_{sw,y,ult}$ – предельные сосредоточенная сила и изгибающие моменты, которые могут быть восприняты арматурой в расчетном поперечном сечении.'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = 'Сосредоточенные моменты от внешней нагрузки $M_{x}$ и $M_{y}$, учитываемые при расчете на продавливание, определяются в соответствии указаниями с п. 8.1.46.'
        string += ' В общем случае, при расположении сосредоточенной силы $F$ внецентренно относительно центра тяжести контура расчетного поперечного сечения'
        string += ' с эксцентриситетами $e_x$ и $e_y$ по формуле:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        if result['M_abs']:
            if result['delta_M_exc']:
                string = '$M_x = (|M_{x,loc}| +  F \\cdot | e_x |) \\cdot \\delta_{Mx}; \\quad'
                string += ' M_y = (|M_{y,loc}| +  F \\cdot | e_y |) \\cdot \\delta_{My}.$'
            else:
                string = '$M_x = |M_{x,loc}| \\cdot \\delta_{Mx} +  F \\cdot | e_x| ; \\quad'
                string += ' M_y =|M_{y,loc}| \\cdot \\delta_{My} +  F \\cdot | e_y|.$'
        else:
            if result['F_dir'] == 'вверх': znak = '+'
            else: znak = '-'
            if result['delta_M_exc']:
                string = '$M_x = |M_{x,loc}' + znak + 'F \\cdot e_x| \\cdot \\delta_{Mx}; \\quad'
                string += ' M_y = |M_{y,loc} ' + znak + ' F \\cdot e_y| \\cdot \\delta_{My}.$'
            else:
                string = '$M_x = |M_{x,loc} \\cdot \\delta_{Mx} ' + znak + ' F \\cdot e_x| ; \\quad'
                string += ' M_y = |M_{y,loc} \\cdot \\delta_{My} ' + znak + ' F \\cdot e_y| .$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'В данных формулах $M_{loc}$ – действующий момент в месте приложения сосредоточенной нагрузки.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Коэффициенты $\\delta_{Mx} \\le 1$ и $\\delta_{My} \\le 1$ учитывают положения п. 8.1.46:'
        string += ' при действии момента в месте приложения сосредоточенной нагрузки половину этого момента учитывают при расчете на продавливание, а другую половину – при расчете по нормальным сечениям.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Предельная продавливающая сила, воспринимаемая бетоном $F_{b,ult}$, вычисляется по формуле (8.88) с учетом формулы (8.89):'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = '$F_{b,ult} = R_{bt} \\cdot h_0 \\cdot u.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Здесь $u$ и $h_0$ – периметр контура расчетного поперечного сечения и приведенная рабочая высота сечения, вычисляемая по формуле:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$h_0 = 0.5 \\cdot (h_{0x} + h_{0y}),$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'где $h_{0x}$ и $h_{0y}$ – рабочая высота сечения для продольной арматуры, расположенной в направлении осей $x$ и $y$.'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = 'Предельная продавливающая сила, воспринимаемая арматурой $F_{sw,ult}$, при равномерном расположении арматуры вдоль контура расчетного поперечного сечения вычисляется по формуле (8.91):'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = '$F_{sw,ult} = 0.8 \\cdot q_{sw} \\cdot u,$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'где $q_{sw}$ – усилие в поперечной арматуре на единицу длины контура расчетного поперечного сечения,'
        string += ' расположенной в пределах расстояния $0.5 \\cdot h_0$ по обе стороны от контура расчетного сечения'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = '$q_{sw} = \\dfrac{R_{sw} \\cdot A_{sw} }{s_w};$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$A_{sw}$ – площадь сечения всей поперечной арматуры расположенная в пределах расстояния $0.5 \\cdot h_0$ по обе'
        string += ' стороны от контура расчетного поперечного сечения по периметру контура расчетного поперечного сечения (площадь поперечной арматуры в одном каркасе);'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = '$s_w$ – шаг поперечной арматуры (шаг каркасов);'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = '$u$ – периметр контура расчетного поперечного сечения.'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = 'Предельный изгибающий момент, воспринимаемый бетоном $M_{b,ult}$, вычисляется по формуле (8.94):'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = '$M_{b,ult} = R_{bt} \\cdot W_b \\cdot h_0.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = 'где $W_b$ – момент сопротивления расчетного поперечного сечения, определяемый в соответствии с п. 8.1.51.'
        st.write(string)
        p = doc.add_paragraph()
        add_text_latex(p, string)
        p.paragraph_format.first_line_indent = Mm(0)

        string = 'Усилия $M_{sw,x,ult}$ и $M_{sw,y,ult}$, воспринимаемые поперечной арматурой, нормальной к продольной оси элемента и'
        string += ' расположенной равномерно вдоль контура расчетного сечения, определяются при действии изгибающего момента'
        string += ' соответственно в направлении осей $x$ и $y$ по формуле (8.97):'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$M_{sw,ult} = 0.8 \\cdot q_{sw} \\cdot W_{sw}.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'В дальнейших расчетах, в соответствии с п. 8.1.52, предполагаем, что'
        string += ' поперечная арматура расположена равномерно вдоль расчетного контура продавливания в пределах зоны, границы'
        string += ' которой отстоят на расстоянии $h_0/2$ в каждую сторону от контура продавливания бетона.'
        string += ' В этом случае значения моментов сопротивления поперечной арматуры при продавливании $W_{sw,x(y)}$ принимают равными соответствующим значениям $W_{bx}$ и $W_{by}$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Введем следующей коэффициент:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$k_{sw} = \\dfrac{F_{sw,ult}}{F_{b,ult}} = \\dfrac{M_{sw,ult}}{M_{b,ult}} ='
        string += ' \\dfrac{0.8 \\cdot q_{sw}}{R_{bt} \\cdot h_0} = \\dfrac{0.8 \\cdot R_{sw} \\cdot A_{sw}}{R_{bt} \\cdot h_0 \\cdot s_w}.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Данный коэффициент можно интерпретировать, как вклад поперечного армирования в несущую способность. C учетом данного коэффициента и предыдущих выкладок условие прочности (8.96) можно переписать в виде:'
        st.write(string)
        doc.add_paragraph().add_run(string)
        string = '$\\dfrac{F}{F_{b,ult} \\cdot (1+k_{sw})} + \\dfrac{M_x}{M_{bx,ult} \\cdot (1+k_{sw})} + \\dfrac{M_y}{M_{by,ult} \\cdot (1+k_{sw}) }  \\le 1.0.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
    
        string = 'В случае, если поперечная арматура не учитывается в расчете, $k_{sw}=0.0$.'
        string += ' В противном случае, если поперечная арматура учитывается в расчете, на коэффициент $k_{sw}$ накладываются следующие ограничения:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$0.25 \\le k_{sw} \\le 1.0.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Нижняя и верхняя граница $k_{sw}$ связаны со ограничениями, приведенными в п. 8.1.48 и 8.1.50.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Во-первых, поперечную арматуру можно учитывать в расчете при выполнении условия $F_{sw,ult} \\ge 0.25 \\cdot F_{b,ult}$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Во-вторых, усилия, воспринимаемые поперечной арматурой, не могут превышать усилия, воспринимаемые бетоном, т.е.'
        string += ' $F_{b,ult} + F_{sw,ult} \\le 2 \\cdot F_{b,ult}$,'
        string += ' $M_{bx,ult} + M_{sw,x,ult} \\le 2 \\cdot M_{bx,ult}$,'
        string += ' $M_{by,ult} + M_{sw,y,ult} \\le 2 \\cdot M_{by,ult}$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Введем дополнительно коэффициенты $k_{b,F}$, $k_{b,M}$ и $k_b$ следующим образом:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$k_{b,F} = \\dfrac{F}{F_{b,ult}}; \\quad'
        string += ' k_{b,M} = \\dfrac{M_x}{M_{bx,ult}} + \\dfrac{M_y}{M_{by,ult}}; \\quad k_b = k_{b,F} + k_{b,M}$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'В результате, условие прочности примет вид:'
        st.write(string)
        doc.add_paragraph().add_run(string)
        string = '$\\dfrac{k_{b,F} + k_{b,M}}{1+k_{sw}} = \\dfrac{k_{b}}{1+k_{sw}} \\le 1.0.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Здесь $k_{b,F}$, $k_{b,M}$ и $k_b$ – коэффициенты использования прочности бетона расчетного поперечного сечения по силе, моментам и суммарный соответственно.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'В соответствии с п. 8.1.46 при действии сосредоточенных моментов и силы в условиях прочности соотношение между действующими'
        string += ' сосредоточенными моментами $M$, учитываемыми при продавливании, и предельными $M_{ult}$ принимают не более'
        string += ' половины соотношения между действующим сосредоточенным усилием $F$ и предельным $F_{ult}$.'
        string += ' Следовательно, $k_{b,M} \\le 0.5 \\cdot k_{b,F}$. В случае нарушения данного условия значение $k_{b,M}$ принимается равным $k_{b,M}=0.5 \\cdot k_{b,F}$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Условие прочности, в случае $k_{b,M} \\ge 0.5 \\cdot k_{b,F}$ примет вид:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$\\dfrac{1.5 \\cdot k_{b,F}}{1+k_{sw}} \\le 1.0.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Из данного неравенства можно получить условие прочности, справедливое для любого значения моментов:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$k_{b,F} \\le \\dfrac{1+k_{sw}}{1.5} \\Rightarrow F \\le \\dfrac{F_{b,ult} \\cdot (1+k_{sw})}{1.5}.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Аналогично при отсутствии поперечной арматуры ($k_{sw}=0.0$):'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$k_{b,F} \\le \\dfrac{1}{1.5} \\Rightarrow F \\le \\dfrac{F_{b,ult}}{1.5}.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Для подбора армирования используется следующая последовательность.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '1. Выполняется расчет геометрических характеристик контура расчетного поперечного сечения.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '2. Выполняется расчет предельных усилий, воспринимаемых бетоном расчетного поперечного сечения.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '3. Вычисление коэффициентов $k_{b,F}$ и $k_{b,M}$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '4. Проверка условия $k_{b,M} \\le 0.5 \\cdot k_{b,F}$. Если условие не выполняется, коэффициент $k_{b,M}$ принимается равным $k_{b,M} = 0.5 \\cdot k_{b,F}$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '5. Вычисление коэффициента использования прочности бетона расчетного поперечного сечения $ k_{b} = k_{b,F} + k_{b,M}$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '6. В случае, если $k_{b} \\le 1.0 $ прочность обеспечена без установки поперечной арматуры.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '7. В случае, если $k_{b} > 2.0 $ прочность не может быть обеспечена, необходимо увеличение габаритов площадки передачи нагрузки, либо толщины плиты, либо класса бетона.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '8. В случае, если $1.0 < k_{b} \\le 2.0 $ требуется установка поперечной арматуры.'
        string += ' В этом случае требуемое соотношение $A_{sw}/s_w$, с учетом выражений для $k_{sw}$ и $q_{sw}$ определяется по формуле:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = '$\\dfrac{A_{sw}}{s_w} \\ge \\dfrac{R_{bt} \\cdot h_0 \\cdot (k_{b} - 1)}{0.8 \\cdot R_{sw}} \\Rightarrow '
        string += 's_w \\le \\dfrac{0.8 \\cdot R_{sw} \\cdot A_{sw} }{R_{bt} \\cdot h_0 \\cdot (k_{b} - 1)}.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        string = 'Примечание. В случае, если $k_{b} \\le 1.25$, для обеспечения требований условия $F_{sw,ult} \\ge 0.25 \\cdot F_{b,ult}$ '
        string += ' в данных формулах необходимо принимать $k_{b} = 1.25$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

def report_full_geometry(result, doc, fig, fig_Aq=None): #Полный расчет геометрических характеристик
    with st.expander('Геометрические характеристики контура'):
        string = 'Геометрические характеристики контура.'
        st.subheader(string)
        doc.add_heading(string, level=2)

        if fig_Aq:
            table = doc.add_table(1, 2)
            p0 = table.cell(0,0).paragraphs[0]
            p0.add_run().add_picture(fig, height=Mm(90))
            p0.paragraph_format.first_line_indent = Mm(0)
            p0.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p00 = table.cell(0,0).add_paragraph()
            p00.style = 'Рис. 3'
            p00.add_run('Эскиз расчетного контура')
            p1 = table.cell(0,1).paragraphs[0]
            p1.add_run().add_picture(fig_Aq, height=Mm(45))
            p1.paragraph_format.first_line_indent = Mm(0)
            p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p10 = table.cell(0,1).add_paragraph()
            p10.style = 'Рис. 3'
            p10.add_run('Габариты наибольшего основания пирамиды продавливания')
        else:
            p = doc.add_paragraph()
            p.add_run().add_picture(fig, width=Mm(80))
            p.paragraph_format.first_line_indent = Mm(0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p =  doc.add_paragraph()
            p.style = 'Рис. 3'
            p.add_run('Эскиз расчетного контура')

        if True: #Длины участков
            string = 'Геометрические характеристики, такие как статические моменты, осевые моменты инерции, моменты сопротивления для расчетного контура вычисляются в НАПРАВЛЕНИИ соответствующих осей.'
            st.write(string)
            #doc.add_paragraph().add_run(string)
            string = 'Длины участков расчетного контура $L_i$, а также длины их проекций $L_{x,i}$ и $L_{y,i}$ в соответствии с эскизом приведены в таблице ниже.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
    
            contours_lens = {'Li, см':[], 'Lxi, см':[], 'Lyi, см':[]}
            for i in range(len(result['contours'])):
                contours_lens['Li, см'].append(str(result['L_arr'][i]))
                contours_lens['Lxi, см'].append(str(result['Lx_arr'][i]))
                contours_lens['Lyi, см'].append(str(result['Ly_arr'][i]))
            contours_lens = pd.DataFrame.from_dict(contours_lens, orient='index', columns=result['sides'])
            st.dataframe(contours_lens, use_container_width=True)

            table = doc.add_table(contours_lens.shape[0]+1, contours_lens.shape[1]+1)
            table.style = 'Стиль_таблицы'
            for j in range(contours_lens.shape[-1]):
                table.cell(0,j+1).text = contours_lens.columns[j]
                table.cell(0,j+1).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            for i in range(contours_lens.shape[0]):
                for j in range(contours_lens.shape[-1]):
                    table.cell(i+1,j+1).text = str(contours_lens.values[i,j])
                    table.cell(i+1,j+1).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            table.cell(0,0).text = 'Участок'
            table.cell(0,0).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            p = table.cell(1,0).paragraphs[0]
            add_text_latex(p, '$L_i$, см')
            p.paragraph_format.first_line_indent = Mm(0)
            p = table.cell(2,0).paragraphs[0]
            add_text_latex(p, '$L_{x,i}$, см')
            p.paragraph_format.first_line_indent = Mm(0)
            p = table.cell(3,0).paragraphs[0]
            add_text_latex(p, '$L_{y,i}$, см')
            p.paragraph_format.first_line_indent = Mm(0)
        
        if True: #Периметр расчетного контура
            string = 'Вычисляем периметр расчетного контура как сумму длин каждого из участков:'
            st.write(string)
            doc.add_paragraph().add_run(string)
            string = '$u = \\sum_i L_i = ' + str(result['L_arr'][0])
            for i in range(1, len(result['L_arr'])):
                string += ' + ' + str(result['L_arr'][i])
            string += ' = ' + str(result['u']) + ' см.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
    
        if True: #Центры тяжести участков
            string = 'Положения центров тяжести $x_{c,0,i}$ и $y_{c,0,i}$ каждого из участков расчетного контура относительно левого нижнего угла колонны приведены в таблице ниже.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            contours_cent = {'xc0i, см':[], 'yc0i, см':[]}
            for i in range(len(result['contours'])):
                contours_cent['xc0i, см'].append(str(result['xc0_arr'][i]))
                contours_cent['yc0i, см'].append(str(result['yc0_arr'][i]))
            contours_cent = pd.DataFrame.from_dict(contours_cent, orient='index', columns=result['sides'])
            st.dataframe(contours_cent, use_container_width=True)
            table = doc.add_table(contours_cent.shape[0]+1, contours_cent.shape[1]+1)
            table.style = 'Стиль_таблицы'
            for j in range(contours_cent.shape[-1]):
                table.cell(0,j+1).text = contours_cent.columns[j]
                table.cell(0,j+1).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            for i in range(contours_cent.shape[0]):
                for j in range(contours_cent.shape[-1]):
                    table.cell(i+1,j+1).text = str(contours_cent.values[i,j])
                    table.cell(i+1,j+1).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            table.cell(0,0).text = 'Участок'
            table.cell(0,0).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            p = table.cell(1,0).paragraphs[0]
            add_math(p, 'x_{c,0,i}, см')
            p.add_run(' ')
            p.paragraph_format.first_line_indent = Mm(0)
            p = table.cell(2,0).paragraphs[0]
            add_math(p, 'y_{c,0,i}, см')
            p.add_run(' ')
            p.paragraph_format.first_line_indent = Mm(0)

        if True: #Статические моменты инерции
            string = 'Статические моменты $S_{bx,0,i}$ и $S_{by,0,i}$ каждого из участков расчетного контура относительно левого нижнего угла колонны вычисляются по формулам:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$S_{bx,0,i} = L_{i} \\cdot x_{c,0,i} ; \\quad S_{by,0,i} = L_{i} \\cdot y_{c,0,i}.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = 'Результаты расчета статических моментов участков расчетного контура по указанным выше приведены в таблице ниже.'
            st.write(string)
            doc.add_paragraph().add_run(string)
            contours_S = {'Sbx0i, см2':[], 'Sby0i, см2':[]}
            for i in range(len(result['contours'])):
                contours_S['Sbx0i, см2'].append(str(result['Sx0_arr'][i]))
                contours_S['Sby0i, см2'].append(str(result['Sy0_arr'][i]))
            contours_S = pd.DataFrame.from_dict(contours_S, orient='index', columns=result['sides'])
            st.dataframe(contours_S, use_container_width=True)
            table = doc.add_table(contours_S.shape[0]+1, contours_S.shape[1]+1)
            table.style = 'Стиль_таблицы'
            for j in range(contours_S.shape[-1]):
                table.cell(0,j+1).text = contours_S.columns[j]
                table.cell(0,j+1).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            for i in range(contours_S.shape[0]):
                for j in range(contours_S.shape[-1]):
                    table.cell(i+1,j+1).text = str(contours_S.values[i,j])
                    table.cell(i+1,j+1).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            table.cell(0,0).text = 'Участок'
            table.cell(0,0).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            p = table.cell(1,0).paragraphs[0]
            add_math(p, 'S_{bx,0,i}, см^2')
            p.add_run(' ')
            p.paragraph_format.first_line_indent = Mm(0)
            #p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p = table.cell(2,0).paragraphs[0]
            add_math(p, 'S_{by,0,i}, см^2')
            p.add_run(' ')
            p.paragraph_format.first_line_indent = Mm(0)

        if True: #Статический момент инерции всего сечения
            string = 'Статические моменты инерции всего расчетного контура относительно левого нижнего угла колонны вычисляем как сумму статических моментов инерции каждого из участков.'
            st.write(string)
            doc.add_paragraph().add_run(string)
            string = '$ S_{bx,0} = \\sum_i S_{bx,0,i}= '
            string += str(result['Sx0_arr'][0])
            for i in range(1, len(result['Sx0_arr'])):
                if result['Sx0_arr'][i] > 0:
                    string += '+'
                string += str(result['Sx0_arr'][i])
            string += '=' + str(result['Sx0']) + ' см^2.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

            string = '$ S_{by,0} = \\sum_i S_{by,0,i}= '
            string += str(result['Sy0_arr'][0])
            for i in range(1, len(result['Sy0_arr'])):
                if result['Sy0_arr'][i] > 0:
                    string += '+'
                string += str(result['Sy0_arr'][i])
            string += ' = ' + str(result['Sy0']) + ' см^2.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

        if True: #Вычисление центра тяжести
            string = 'Положение геометрического центра тяжести контура относительно левого нижнего угла колонны.'
            st.write(string)
            doc.add_paragraph().add_run(string)
            string = '$ x_c = \\dfrac{S_{bx,0}}{u} = '
            string += '\\dfrac{' + str(result['Sx0']) + '}{' + str(result['u']) + '}='
            string += str(result['xc']) + ' см.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

            string = '$ y_c = \\dfrac{S_{by,0}}{u} = '
            string += '\\dfrac{' + str(result['Sy0']) + '}{' + str(result['u']) + '}='
            string += str(result['yc']) + ' см.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

        if True: #Эксцентриситет продольной силы
            if result["ex"] !=0 or result["ey"] !=0:
                string = 'Эксцентриситет центра тяжести расчетного контура относительно центра колонны.'
                st.write(string)
                doc.add_paragraph().add_run(string)
                if result["ex"] != 0:
                    string = 'Вдоль оси $x$:'
                    st.write(string)
                    add_text_latex(doc.add_paragraph(), string)
                    string = '$e_x = x_c - b/2= ' + str(result['xc']) + ' - ' +  str(round(result['b']/2,2))
                    string += ' = ' + str(result["ex"]) + ' см.$'
                    st.write(string)
                    add_text_latex(doc.add_paragraph(), string)
                if result["ey"] != 0:
                    string = 'Вдоль оси $y$:'
                    st.write(string)
                    add_text_latex(doc.add_paragraph(), string)
                    string = '$e_y = y_c - h/2= ' + str(result['yc']) + ' - ' +  str(round(result['h']/2,2))
                    string += ' = ' + str(result["ey"]) + ' см.$'
                    st.write(string)
                    add_text_latex(doc.add_paragraph(), string)

        if True: #Вычисление центров тяжести участков
            string = 'Вычисляем координаты центров тяжести каждого из элементов расчетного контура относительно центра тяжести всего расчетного контура по формулам:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$x_{c,i} = x_{c,0,i} - x_c; \\quad y_{c,i} = y_{c,0,i} - y_c.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = 'Результаты расчета по указанным выше формулам приведены в таблице ниже.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            contours_cent = {'xci, см':[], 'yci, см':[]}
            for i in range(len(result['contours'])):
                contours_cent['xci, см'].append(str(result['xc_arr'][i]))
                contours_cent['yci, см'].append(str(result['yc_arr'][i]))
            contours_cent = pd.DataFrame.from_dict(contours_cent, orient='index', columns=result['sides'])
            st.dataframe(contours_cent, use_container_width=True)
            table = doc.add_table(contours_cent.shape[0]+1, contours_cent.shape[1]+1)
            table.style = 'Стиль_таблицы'
            for j in range(contours_cent.shape[-1]):
                table.cell(0,j+1).text = contours_cent.columns[j]
                table.cell(0,j+1).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            for i in range(contours_cent.shape[0]):
                for j in range(contours_cent.shape[-1]):
                    table.cell(i+1,j+1).text = str(contours_cent.values[i,j])
                    table.cell(i+1,j+1).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            table.cell(0,0).text = 'Участок'
            table.cell(0,0).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            p = table.cell(1,0).paragraphs[0]
            add_math(p, 'x_{c,i}, см')
            p.add_run(' ')
            p.paragraph_format.first_line_indent = Mm(0)
            p = table.cell(2,0).paragraphs[0]
            add_math(p, 'y_{c,i}, см')
            p.add_run(' ')
            p.paragraph_format.first_line_indent = Mm(0)

        if True: #Вычисление наиболее удаленных точек
            string = 'Расстояние до наиболее удаленных от геометрического центра тяжести точек (левой, $x_L$; правой, $x_R$; нижней, $y_B$; верхней, $y_T$) расчетного контура приведено в таблице ниже:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            max_values = [result['xL'], result['xR'], result['yB'], result['yT']]
            max_points = {'Расстояние, см': max_values}
            max_points = pd.DataFrame.from_dict(max_points, orient='index', columns=['Левая','Правая','Нижняя','Верхняя'])
            st.dataframe(max_points, use_container_width=True)
            table = doc.add_table(2, 5)
            table.style = 'Стиль_таблицы'
            table.cell(1,0).text = 'Расстояние, см'
            table.cell(1,0).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            p = table.cell(0,1).paragraphs[0]
            add_math(p, 'x_{L}')
            p.paragraph_format.first_line_indent = Mm(0)
            p = table.cell(0,2).paragraphs[0]
            add_math(p, 'x_{R}')
            p.paragraph_format.first_line_indent = Mm(0)
            p = table.cell(0,3).paragraphs[0]
            add_math(p, 'y_{B}')
            p.paragraph_format.first_line_indent = Mm(0)
            p = table.cell(0,4).paragraphs[0]
            add_math(p, 'y_{T}')
            p.paragraph_format.first_line_indent = Mm(0)
            for i in range(len(max_values)):
                p = table.cell(1,i+1).paragraphs[0]
                add_text_latex(p, str(max_values[i]))
                p.paragraph_format.first_line_indent = Mm(0)

            string = 'Расстояние до наиболее удаленных от центра тяжести точек расчетного контура составляет:'
            st.write(string)
            doc.add_paragraph().add_run(string)
            string = '$x_{\\max} = \\max(|x_{L}|,x_{R})='
            string += str(result['xmax'])
            string += 'см; \\quad y_{\\max} = \\max(|y_{B}|,y_{T})='
            string += str(result['ymax'])
            string += 'см.$'
            st.write(string)
            add_math(doc.add_paragraph(), string.replace('$',''))

        if True: #Вычисление собственных моментов инерции участков
            string = 'Собственные моменты инерции участков расчетного контура вычисляются по формулам:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$I_{bx,0,i} = \\dfrac{L_{x,i}^3}{12}; \\quad I_{by,0,i} = \\dfrac{L_{y,i}^3}{12}.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = 'В таблице ниже приведены результаты расчета по указанным выше формулам.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            contours_I0 = {'Ibx0i, см3':[], 'Iby0i, см3':[]}
            for i in range(len(result['contours'])):
                contours_I0['Ibx0i, см3'].append(str(result["Ix0_arr"][i]))
                contours_I0['Iby0i, см3'].append(str(result["Iy0_arr"][i]))
            contours_I0 = pd.DataFrame.from_dict(contours_I0, orient='index', columns=result['sides'])
            st.dataframe(contours_I0, use_container_width=True)
            table = doc.add_table(contours_I0.shape[0]+1, contours_I0.shape[1]+1)
            table.style = 'Стиль_таблицы'
            for j in range(contours_I0.shape[-1]):
                table.cell(0,j+1).text = contours_I0.columns[j]
                table.cell(0,j+1).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            for i in range(contours_I0.shape[0]):
                for j in range(contours_I0.shape[-1]):
                    table.cell(i+1,j+1).text = str(contours_I0.values[i,j])
                    table.cell(i+1,j+1).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            table.cell(0,0).text = 'Участок'
            table.cell(0,0).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            p = table.cell(1,0).paragraphs[0]
            add_math(p, 'I_{bx,0,i}, см^3')
            p.add_run(' ')
            p.paragraph_format.first_line_indent = Mm(0)
            #p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p = table.cell(2,0).paragraphs[0]
            add_math(p, 'I_{by,0,i}, см^3')
            p.add_run(' ')
            p.paragraph_format.first_line_indent = Mm(0)

        if True: #Вычисление моментов инерции участков
            string = 'Для вычисления моментов инерции участков контура относительно геометрического центра всего контура используются формулы переноса осей, приведенные ниже.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$I_{bx,i} = I_{bx,0,i} +  L_i \\cdot x_{c,i}^2; \\quad I_{by,i} = I_{by,0,i} +  L_i \\cdot y_{c,i}^2.$'
            st.write(string)
            add_math(doc.add_paragraph(), string.replace('$',''))
            string = 'Результаты расчета моментов инерции участков контура относительно геометрического центра всего контура по указанным выше формулам приведены в таблице ниже.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            contours_I = {'Ibxi, см3':[], 'Ibyi, см3':[]}
            contours_I_names = []
            for i in range(len(result['contours'])):
                contours_I['Ibxi, см3'].append(str(result['Ix_arr'][i]))
                contours_I['Ibyi, см3'].append(str(result['Iy_arr'][i]))
            contours_I = pd.DataFrame.from_dict(contours_I, orient='index', columns=result['sides'])
            st.dataframe(contours_I, use_container_width=True)
            table = doc.add_table(contours_I.shape[0]+1, contours_I.shape[1]+1)
            table.style = 'Стиль_таблицы'
            for j in range(contours_I.shape[-1]):
                table.cell(0,j+1).text = contours_I.columns[j]
                table.cell(0,j+1).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            for i in range(contours_I.shape[0]):
                for j in range(contours_I.shape[-1]):
                    table.cell(i+1,j+1).text = str(contours_I.values[i,j])
                    table.cell(i+1,j+1).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            table.cell(0,0).text = 'Участок'
            table.cell(0,0).paragraphs[0].paragraph_format.first_line_indent = Mm(0)
            p = table.cell(1,0).paragraphs[0]
            add_math(p, 'I_{bx,i}, см^3')
            p.add_run(' ')
            p.paragraph_format.first_line_indent = Mm(0)
            #p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p = table.cell(2,0).paragraphs[0]
            add_math(p, 'I_{by,i}, см^3')
            p.add_run(' ')
            p.paragraph_format.first_line_indent = Mm(0)
            #for row in table.rows:
            #    row._element.attrib['keepTogether'] = '1'
            #    row.allow_break_between_rows = False
            #    row.allow_break_between_pages = False

        if True: #Моменты инерции всего сечения
            string = 'Моменты инерции всего расчетного контура вычисляем как сумму моментов инерции каждого из участков.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$ I_{bx} = \\sum_i{I_{bx,i}}= '
            string += str(result['Ix_arr'][0])
            for i in range(1, len(result['Ix_arr'])):
                if result['Ix_arr'][i] > 0:
                    string += '+'
                string += str(result['Ix_arr'][i])
            string += '=' + str(result['Ix']) + ' см^3;$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

            string = '$ I_{by} = \\sum_i{I_{by,i}}= '
            string += str(result['Iy_arr'][0])
            for i in range(1, len(result['Iy_arr'])):
                if result['Iy_arr'][i] > 0:
                    string += '+'
                string += str(result['Iy_arr'][i])
            string += '=' + str(result['Iy']) + ' см^3.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

        if True: #Моменты сопротивления расчетного контура
            string = 'Моменты сопротивления расчетного контура вычисляются по формулам (8.98):'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$ W_{bx} = \\dfrac{I_{bx}}{x_{\\max}} = \\dfrac{'
            string += str(result['Ix']) +'}{'
            string += str(result['xmax']) + '} = ' + str(result['Wx'])
            string += ' см^2;$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

            string = '$ W_{by} = \\dfrac{I_{by}}{y_{\\max}} = \\dfrac{'
            string += str(result['Iy']) + '}{'
            string += str(result['ymax']) + '} = ' +str(result['Wy'])
            string += ' см^2.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

def report_short_geometry(result, doc, fig, fig_Aq=None): #Краткий расчет геометрических характеристик
     with st.expander('Геометрические характеристики контура'):
        string = 'Геометрические характеристики контура.'
        st.subheader(string)
        doc.add_heading(string, level=2)
        string = 'Геометрические характеристики, такие как осевые моменты инерции и моменты сопротивления для расчетного контура вычисляются в НАПРАВЛЕНИИ соответствующих осей.'
        st.write(string)
        if fig_Aq:
            table = doc.add_table(1, 2)
            p0 = table.cell(0,0).paragraphs[0]
            p0.add_run().add_picture(fig, height=Mm(90))
            p0.paragraph_format.first_line_indent = Mm(0)
            p0.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p00 = table.cell(0,0).add_paragraph()
            p00.style = 'Рис. 3'
            p00.add_run('Эскиз расчетного контура')
            p1 = table.cell(0,1).paragraphs[0]
            p1.add_run().add_picture(fig_Aq, height=Mm(45))
            p1.paragraph_format.first_line_indent = Mm(0)
            p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p10 = table.cell(0,1).add_paragraph()
            p10.style = 'Рис. 3'
            p10.add_run('Габариты наибольшего основания пирамиды продавливания')
        else:
            p = doc.add_paragraph()
            p.add_run().add_picture(fig, width=Mm(80))
            p.paragraph_format.first_line_indent = Mm(0)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p =  doc.add_paragraph()
            p.style = 'Рис. 3'
            p.add_run('Эскиз расчетного контура')

        if True: #Периметр расчетного контура
            #string = 'Геометрические характеристики, такие как осевые моменты инерции и моменты сопротивления для расчетного контура вычисляются в НАПРАВЛЕНИИ соответствующих осей.'
            #st.write(string)
            #add_text_latex(doc.add_paragraph(), string)
            string = 'Периметр расчетного контура:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$u = ' + str(result['u']) + ' см.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

        if True: #Центра тяжести
            string = 'Координаты центра тяжести контура относительно левого нижнего угла колонны:'
            st.write(string)
            doc.add_paragraph().add_run(string)
            string = '$ x_c = ' + str(result['xc']) + ' см; \\quad  y_c = ' + str(result['yc']) + ' см.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
        
        if True: #Эксцентриситет продольной силы
            if result["ex"] !=0 or result["ey"] !=0:
                string = 'Эксцентриситеты продольной силы:'
                st.write(string)
                add_text_latex(doc.add_paragraph(), string)
                string = '$'
                if result["ex"] !=0:
                    string += 'e_x = ' + str(result["ex"]) +  ' см'
                if result["ey"] !=0:
                    if result["ex"] !=0:
                        string += '; \\quad '
                    string += 'e_y = ' + str(result["ey"]) +  ' см'
                string += '.$'
                st.write(string)
                add_text_latex(doc.add_paragraph(), string)

        if True: #Вычисление наиболее удаленных точек
            string = 'Расстояние до наиболее удаленных от центра тяжести точек расчетного контура составляет:'
            st.write(string)
            doc.add_paragraph().add_run(string)
            string = '$x_{\\max} = ' + str(result["xmax"]) + ' см; \\quad '
            string += ' y_{\\max} =' + str(result["ymax"]) + ' см.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

        if True: #Моменты инерции всего сечения
            string = 'Моменты инерции расчетного контура:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$ I_{bx} = ' + str(result["Ix"]) + ' см^3; \\quad '
            string += ' I_{by} = '  + str(result["Iy"]) + ' см^3.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

        if True: #Моменты сопротивления расчетного контура
            string = 'Моменты сопротивления расчетного контура:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$ W_{bx} = ' + str(result["Wx"]) + ' см^2; \\quad'
            string += ' W_{by} = ' + str(result["Wy"]) + ' см^2.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

def report_concrete_ultimate_forces(result, doc): #Расчет предельных усилий, воспринимаемых бетоном
    with st.expander('Предельные усилия, воспринимаемые бетоном'):
        if True: #Предельная продавливающая сила, воспринимаемая бетоном
            string = 'Предельные усилия, воспринимаемые бетоном.'
            st.subheader(string)
            doc.add_heading(string, level=2)
            string = 'Предельная сила, воспринимаемая бетоном $F_{b,ult}$, вычисляется по формуле (8.88) с учетом формулы (8.89):'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$F_{b,ult} = R_{bt} \\cdot h_0 \\cdot u = '
            string += str(result['Rbt']) + ' \\cdot ' +  str(result['h0']) + ' \\cdot ' + str(result['u']) + ' = ' + str(result['Fbult']) + ' тс.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

            string = 'Предельно допустимое значение силы (с учетом положений п. 8.1.46), при которой допускается не учитывать изгибающие моменты:'
            st.write(string)
            doc.add_paragraph().add_run(string)
            string = '$F_{b,ult}/1.5 = ' + str(result['Fbult']) + '/1.5 = ' +  str(round(result['Fbult']/1.5, 1)) + ' тс.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

        if True: #Предельные моменты, воспринимаемые бетонным сечением
            string = 'Предельные моменты, воспринимаемые бетоном в расчетном поперечном сечении, вычисляются по формулам (8.94):'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$ M_{bx,ult} = R_{bt} \\cdot h_0 \\cdot W_{bx} = '
            string += str(result['Rbt']) + ' \\cdot ' + str(result['h0']) + ' \\cdot ' + str(result['Wx']) + '/100 = ' + str(result['Mbxult']) + ' тсм;$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

            string = '$ M_{by,ult} = R_{bt} \\cdot h_0 \\cdot W_{by} = '
            string += str(result['Rbt']) + ' \\cdot ' + str(result['h0']) + ' \\cdot ' + str(result['Wy']) + '/100 = ' + str(result['Mbyult']) + ' тсм.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

            string = 'Примечание: деление на 100 для перевода см в м.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

def report_acting_forces(result, doc): #Расчетные усилия
    with st.expander('Усилия, учитываемых в расчете'):
        string = 'Усилия, учитываемые в расчете.'
        st.subheader(string)
        doc.add_heading(string, level=2)
        string = 'Сосредоточенная сила, направленная ' + str(result['F_dir']) + ' $F=' + str(result['F0']) + ' тс$.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        if round(abs(result['Fq']), 1) != 0.0:
            string = 'Разгружающая сила $F_q = q \\cdot A_q = ' + str(result['q']) + '\\cdot' + str(result['Aq']) + '=' + str(result['Fq']) + ' тс$.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = 'В расчете принимаем $F = F - F_q = ' + str(result['F0']) + '-' + str(result['Fq']) + '=' + str(result['F']) + ' тс$.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

        string = 'Сосредоточенные моменты.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

        if result["ex"] !=0: #Если есть эксцентриситет вдоль х
            if result['M_abs']: #Если считаем момент от эксцентриситета всегда догружающим
                if result['delta_M_exc']: #Если учитываем дельта к эксцентриситету
                    string = '$M_x = (|M_{x,loc}| + F \\cdot |e_x|) \\cdot \\delta_{Mx} = '
                    string += '(|' + str(result['Mxloc']) + '| + ' + str(result['F']) + '\\cdot |' + str(result['ex']) + '/100|) \\cdot ' + str(result['deltaMx']) +  ' =  ' + str(result['Mx']) 
                else:  #Если не учитываем дельта к эксцентриситету
                    string = '$M_x = |M_{x,loc}| \\cdot \\delta_{Mx} + F \\cdot |e_x|  = '
                    string += '|' + str(result['Mxloc']) + '| \\cdot'  + str(result['deltaMx']) + ' + ' + str(result['F']) + '\\cdot |' + str(result['ex']) + '/100| '  ' =  ' + str(result['Mx']) 
            else: #Если учитываем знаки моментов
                if result['F_dir'] == 'вверх': znak = '+'
                if result['F_dir'] == 'вниз': znak = '-'
                if result['delta_M_exc']: #Если учитываем дельта к эксцентриситету
                    string = '$M_x = |M_{x,loc}' + znak +  'F \\cdot e_x| \\cdot \\delta_{Mx} = '
                    string += '|' + str(result['Mxloc']) + znak + str(result['F']) + '\\cdot ' 
                    if result['ex'] <0: string += '(' +  str(result['ex'])  + ')'
                    else: string += str(result['ex'])
                    string += '/100| \\cdot ' + str(result['deltaMx']) +  ' =  ' + str(result['Mx']) 
                else:  #Если не учитываем дельта к эксцентриситету
                    string = '$M_x = |M_{x,loc} \\cdot \\delta_{Mx}' + znak +  'F \\cdot e_x|  = '
                    string += '|' + str(result['Mxloc']) + ' \\cdot'  + str(result['deltaMx']) + znak + str(result['F']) + '\\cdot '
                    if result['ex'] <0: string += '(' +  str(result['ex'])  + ')'
                    else: string += str(result['ex'])
                    string +=  '/100| '  ' =  ' + str(result['Mx']) 
            string +=  ' тсм.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

        if result["ey"] !=0: #Если есть эксцентриситет вдоль х
            if result['M_abs']: #Если считаем момент от эксцентриситета всегда догружающим
                if result['delta_M_exc']: #Если учитываем дельта к эксцентриситету
                    string = '$M_y = (|M_{y,loc}| + F \\cdot |e_y|) \\cdot \\delta_{My} = '
                    string += '(|' + str(result['Myloc']) + '| + ' + str(result['F']) + '\\cdot |' + str(result['ey']) + '/100|) \\cdot ' + str(result['deltaMy']) +  ' =  ' + str(result['My']) 
                else:  #Если не учитываем дельта к эксцентриситету
                    string = '$M_y = |M_{y,loc}| \\cdot \\delta_{My} + F \\cdot |e_y|  = '
                    string += '|' + str(result['Myloc']) + '| \\cdot'  + str(result['deltaMy']) + ' + ' + str(result['F']) + '\\cdot |' + str(result['ey']) + '/100| '  ' =  ' + str(result['My']) 
            else: #Если учитываем знаки моментов
                if result['F_dir'] == 'вверх': znak = '+'
                if result['F_dir'] == 'вниз': znak = '-'
                if result['delta_M_exc']: #Если учитываем дельта к эксцентриситету
                    string = '$M_y = |M_{y,loc}' + znak +  'F \\cdot e_y| \\cdot \\delta_{My} = '
                    string += '|' + str(result['Myloc']) + znak + str(result['F']) + '\\cdot ' 
                    if result['ey'] <0: string += '(' +  str(result['ey'])  + ')'
                    else: string += str(result['ey'])
                    string += '/100| \\cdot ' + str(result['deltaMy']) +  ' =  ' + str(result['My']) 
                else:  #Если не учитываем дельта к эксцентриситету
                    string = '$M_y = |M_{y,loc} \\cdot \\delta_{My}' + znak +  'F \\cdot e_y|  = '
                    string += '|' + str(result['Myloc']) + ' \\cdot'  + str(result['deltaMy']) + znak + str(result['F']) + '\\cdot '
                    if result['ey'] <0: string += '(' +  str(result['ey'])  + ')'
                    else: string += str(result['ey'])
                    string +=  '/100| '  ' =  ' + str(result['My']) 
            string +=  ' тсм.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
    
        if result["ex"] !=0 or result["ey"] !=0:
            string = 'Примечание: деление на 100 для перевода см в м.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

        if result["ex"] == 0:
            string = '$M_x = |M_{x,loc}| \\cdot \\delta_{Mx} = | ' + str(result['Mxloc']) +'| \\cdot ' + str(result['deltaMx']) + '=' + str(result['Mx']) +  ' тсм.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

        if result["ey"] == 0:
            string = '$M_y = |M_{y,loc}| \\cdot \\delta_{My} = | ' + str(result['Myloc']) +'| \\cdot ' + str(result['deltaMy']) + '=' + str(result['My']) +  ' тсм.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)

def report_concrete_strength (result, doc): #Прочность по бетону
    with st.expander('Проверка прочности по бетону'):
        string = 'Проверка прочности бетона расчетного поперечного сечения.'
        st.subheader(string)
        doc.add_heading(string, level=2)
        if True: #Проверка по продольной силе
            string = 'Коэффициент использования прочности бетона расчетного поперечного сечения по силе:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$k_{b,F}=\\dfrac{F}{F_{b,ult}}=\\dfrac{'
            string += str(result['F']) + '}{' + str(result['Fbult']) + '} = ' + str(result['kbF']) + '.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
        if True: #Проверка по моментам
            string = 'Коэффициент использования прочности бетона по моментам:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$k_{b,M}=\\dfrac{M_x}{M_{bx,ult}} + \\dfrac{M_y}{M_{by,ult}} =\\dfrac{'
            string += str(result['Mx']) + '}{' + str(result['Mbxult']) + '} +  \\dfrac{' + str(result['My']) + '}{' + str(result['Mbyult'])
            string += '}=' + str(result['kbM0']) + '.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            if result['kbM0'] != result['kbM']:
                string = 'Условие $k_{b,M} \\le 0.5 \\cdot k_{b,F}$ не выполняется. Вклад моментов ограничивается в соответствии с указаниями п. 8.1.46.'
                string += ' В расчете принимаем:'
                st.write(string)
                add_text_latex(doc.add_paragraph(), string)
                string = '$k_{b,M} = 0.5 \\cdot k_{b,F} =' + str(result['kbM']) + '.$'
                st.write(string)
                add_text_latex(doc.add_paragraph(), string)

        if True: #Суммарно
            string = 'Суммарный коэффициент использования прочности бетона:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$k_b = k_{b,F} + k_{b,M}=' + str(result['kbF']) + ' + ' + str(result['kbM']) + ' = ' + str(result['kb']) + ' .$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
        
        if True: #Результаты проверки прочности по бетону
            if result['kb'] <= 1:
                string = 'Так как $k_{b}=' + str(result['kb'])+ '\\le 1$ прочность обеспечена без установки поперечной арматуры.'
                st.write(string)
                add_text_latex(doc.add_paragraph(), string)
            if result['kb'] > 2:
                string = 'Так как $k_{b}=' + str(result['kb'])+ '>2$ прочность не может быть обеспечена, необходимо увеличение габаритов площадки передачи нагрузки, либо толщины плиты, либо класса бетона.'
                st.write(string)
                add_text_latex(doc.add_paragraph(), string)
            if 1 < result['kb'] <= 2:
                string = 'Так как $1< k_{b}=' + str(result['kb'])+  ' \\le 2$ требуется установка поперечной арматуры.'
                st.write(string)
                add_text_latex(doc.add_paragraph(), string)

def report_reinf_ultimate_forces(result, doc): #Расчет предельных усилий, воспринимаемых арматурой
    with st.expander('Предельные усилия, воспринимаемые арматурой'):
        string = 'Предельные усилия, воспринимаемые поперечной арматурой.'
        st.subheader(string)
        doc.add_heading(string, level=2)
        string = 'Предельная продавливающая сила, воспринимаемая поперечной арматурой $F_{sw,ult}$, вычисляется по формуле (8.91):'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = '$F_{sw,ult} = 0.8 \\cdot q_{sw} \\cdot u = '
        string += '0.8 \\cdot' + str(result['qsw']) + '\\cdot' + str(result['u']) + '='
        string += str(result['Fswult0']) + ' тс.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        if result['Fsw_code'] == 2: #Если поперечную арматуру нельзя учитывать в расчете
            string = 'Так как $F_{sw,ult} \\le 0.25 \\cdot F_{b,ult} = ' + str(round(0.25 * result['Fbult'],1)) + '$ , поперечную арматуру не учитываем в расчете. В расчете принимаем: '
            string += '$F_{sw,ult} = M_{sw,x,ult} = M_{sw,y,ult} = 0.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
        string = ''
        if result['Fsw_code'] != 2: #Если поперечную арматуру можно учитывать в расчете
            string += 'Условие $F_{sw,ult} ' +  ' \\ge 0.25 \\cdot F_{b,ult} =' + str(round(0.25*result['Fbult'],1)) +  ' $ выполняется.'
            string += ' Поперечная арматура учитывается в расчете. '
        if result['Fsw_code'] == 1: #Если несущая способность по поперечной арматуре превышает несущую способность по бетону
            string += 'Условие $F_{sw,ult} \\le F_{b,ult}$ не выполняется, вклад поперечного армирования ограничиваем. В расчете принимаем: '
            string += '$F_{sw,ult} = F_{b,ult}=' + str(result['Fswult']) +  ' тс.$'
        if result['Fsw_code'] == 0: #Если несущая способность по поперечной арматуре не превышает несущую способность по бетону
            string += 'Условие $F_{sw,ult} \\le F_{b,ult}$ выполняется, вклад поперечного армирования не ограничиваем.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        if result['Fsw_code'] != 2: #Если поперечную арматуру можно учитывать в расчете
            string = 'Предельные моменты, воспринимаемые поперечной арматурой, вычисляются по формуле (8.97):'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$M_{sw,ult}= 0.8 \\cdot q_{sw} \\cdot W_{sw}.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            
            string = 'В соответствии с п. 8.1.52, принято, что поперечная арматура расположена равномерно вдоль расчетного контура продавливания, т.е. $W_{sw}=W_b$. В результате:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$M_{sw,x,ult}= 0.8 \\cdot q_{sw} \\cdot W_{bx} = '
            string += '0.8 \\cdot' + str(result['qsw']) + '\\cdot' + str(result['Wx']) + '/100='
            string += str(result['Mswxult0']) + ' тсм;$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$ M_{sw,y,ult}= 0.8 \\cdot q_{sw} \\cdot W_{by} = '
            string += '0.8 \\cdot' + str(result['qsw']) + '\\cdot' + str(result['Wy']) + '/100='
            string += str(result['Mswyult0']) + ' тсм.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = 'Примечание: деление на 100 для перевода см в м.'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            if result['Mswx_code'] == 0:
                string = 'Условия $M_{sw,x,ult} \\le M_{bx,ult}$ и $M_{sw,y,ult} \\le M_{by,ult}$ выполняются, вклад поперечного армирования не ограничиваем.'
            if result['Mswx_code'] == 1:
                string = 'Условя $M_{sw,x,ult} \\ge M_{bx,ult}$ и $M_{sw,y,ult} \\ge M_{by,ult}$ не выполняются, вклад поперечного армирования ограничиваем. В расчете принимаем:'
                st.write(string)
                add_text_latex(doc.add_paragraph(), string)
                string = '$M_{sw,x,ult} = M_{bx,ult}=' + str(result['Mswxult']) +  ' тсм; \\quad'
                string += ' M_{sw,y,ult} = M_{by,ult}=' + str(result['Mswyult'])
                string += ' тсм.$'
                st.write(string)
                add_text_latex(doc.add_paragraph(), string)

def report_solve_sw_min(result, doc): #Определением минимального шага поперечной арматуры
    with st.expander('Подбор шага поперечного армирования'):
        string = 'Определение максимального шага поперечного армирования.'
        st.subheader(string)
        doc.add_heading(string, level=2)
        string = ''
        if result['kb_sw_code'] == 1:
            string += 'Для обеспечения требований $F_{sw,ult} \\ge 0.25 \\cdot F_{b,ult}$ при подборе максимального шага поперечного армирования принимаем $k_b = 1.25$.'
        string += ' Максимально допустимый шаг поперечного армирования при заданном $A_{sw} = ' + str(result['Asw']) + ' см^2$ из условия прочности составляет:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = '$s_w \\le \\dfrac{0.8 \\cdot R_{sw} \\cdot A_{sw} }{R_{bt} \\cdot h_0 \\cdot (k_{b} - 1)} = '
        string += '\\dfrac{0.8 \\cdot ' + str(result['Rsw']) + ' \\cdot ' + str(result['Asw']) + ' }{' + str(result['Rbt']) + ' \\cdot ' + str(result['h0']) + ' \\cdot (' + str(result['kb']) + ' - 1)} = '
        string += str(result['sw_min0']) +  ' см.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = 'Согласно конструктивным требованиям, шаг поперечного армирования не должен превышать величин $d_x/4$ и $d_y/4$, где $d_x$ и $d_y$ – габариты расчетного контура.'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = '$d_x/4 = ' + str(result['dx']) + '/ 4 = ' + str(round(result['dx']/4,1))  +  ' см ; \\quad d_y/4 = ' + str(result['dy']) + '/ 4 = ' + str(round(result['dy']/4,1))  + ' см.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = ''
        if result['sw_min_code'] == 1:
            string += 'Подобранный шаг не удовлетворяет конструктивным требованиям. В расчете принимаем:'
        else:
            string += 'Подобранный шаг удовлетворяет конструктивным требованиям. В расчете принимаем:'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)
        string = '$s_w = ' + str(result['sw']) + ' см; q_{sw} = \dfrac{R_{sw} \\cdot A_{sw}}{s_w} = \\dfrac{' + str(result['Rsw']) + ' \\cdot ' + str(result['Asw']) + '}{' + str(result['sw']) + '} = ' + str(result['qsw']) + ' тс/см.$'
        st.write(string)
        add_text_latex(doc.add_paragraph(), string)

def report_full_strength (result, doc): #Прочность по c поперечной арматурой
    with st.expander('Проверка прочности с арматурой'):
        string = 'Проверка прочности расчетного поперечного сечения с поперечной арматурой.'
        st.subheader(string)
        doc.add_heading(string, level=2)
        if True: #Проверка по продольной силе
            string = 'Коэффициент использования прочности расчетного поперечного сечения по силе:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$k_{F}=\\dfrac{F}{F_{b,ult} + F_{sw,ult}}=\\dfrac{'
            string += str(result['F']) + '}{' + str(result['Fbult']) + '+' + str(result['Fswult']) + '} = ' + str(result['kF']) + '.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
        
        if True: #Проверка по моментам
            string = 'Коэффициент использования прочности расчетного поперечного сечения по моментам:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$k_{M}=\\dfrac{M_x}{M_{bx,ult} + M_{sw,x,ult}} + \\dfrac{M_y}{M_{by,ult} + M_{sw,y,ult}} =\\dfrac{'
            string += str(result['Mx']) + '}{' + str(result['Mbxult']) + '+' + str(result['Mswxult']) + '} +  \\dfrac{' + str(result['My']) + '}{' + str(result['Mbyult'])  + '+' + str(result['Mswyult'])
            string += '}=' + str(result['kM0']) + '.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            if result['kM_code'] == 1:
                string = 'Условие $k_{M} \\le 0.5 \\cdot k_{F}$ не выполняется. Вклад моментов ограничивается в соответствии с указаниями п. 8.1.46.'
                string += ' В расчете принимаем:'
                st.write(string)
                add_text_latex(doc.add_paragraph(), string)
                string = '$k_{M} = 0.5 \\cdot k_{F} =' + str(result['kM']) + '.$'
                st.write(string)
                add_text_latex(doc.add_paragraph(), string)

        if True: #Суммарно
            string = 'Суммарный коэффициент использования прочности расчетного поперечного сечения:'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
            string = '$k=k_{F}+k_{M}=' + str(result['kF']) + '+' + str(result['kM']) + '=' + str(result['k']) + '.$'
            st.write(string)
            add_text_latex(doc.add_paragraph(), string)
        if True: #Результаты проверки прочности с арматурой
            if result['k'] <= 1:
                string = 'Так как $k=' + str(result['k'])+ '<1$ прочность обеспечена.'
                st.write(string)
                add_text_latex(doc.add_paragraph(), string)
            if 2 >= result['k'] > 1:
               string = 'Так как $k=' + str(result['k'])+ '>1$ '
               if result['kb'] < 2:
                   string += 'прочность не обеспечена. Требуется увеличение поперечного армирования.'
               if result['kb'] > 2:
                   string += ' и $k_b=' + str(result['kb']) + '>2$, прочность не может быть обеспечена, необходимо увеличение габаритов колонны, толщины плиты или класса бетона.'
               st.write(string)
               add_text_latex(doc.add_paragraph(), string)
